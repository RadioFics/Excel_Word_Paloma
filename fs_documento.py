"""
fs_documento.py — Motor de regiones sobre un documento de Word vivo.

A diferencia de generador_fs.py (que RENDERIZA una plantilla y produce un
.docx nuevo y desechable), este módulo ACTUALIZA EN EL SITIO un documento
que ya existe: localiza las anclas del contrato (fs_contrato.py), reescribe
únicamente su interior y no visita nada más. La prosa que una persona
redacte alrededor sobrevive intacta a cada refresco; si la borra, se queda
borrada, porque el motor nunca la vuelve a inyectar.

Órdenes
-------
    python fs_documento.py construir  <doc.docx> [--desde base.docx]
    python fs_documento.py reparar    <doc.docx>
    python fs_documento.py refrescar  <doc.docx> [libro.xlsx] [--sin-registro]
    python fs_documento.py insertar   <doc.docx> <clave> <campo> [--zona n]
    python fs_documento.py verificar  <doc.docx> [libro.xlsx]
    python fs_documento.py catalogo   [libro.xlsx]
    python fs_documento.py plantilla  <destino.docx>
    python fs_documento.py proteger   <doc.docx> --clave <clave>
    python fs_documento.py desproteger <doc.docx>

'construir' y 'reparar' son la misma operación: añaden lo que falte para
que el documento cumpla el contrato, sin duplicar ni borrar lo que ya haya.
Se pueden correr sobre un documento con meses de redacción encima.

ADVERTENCIA: el refresco escribe sobre el archivo indicado. Si vive en
OneDrive, ciérrelo en Word antes de refrescar (o Word y el motor pelearán
por el archivo). Siempre se deja una copia .bak junto al original.
"""
import sys
import json
import copy
import shutil
import hashlib
import struct
import base64
import os
import traceback
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# El Python portable de .\python\ es la distribución "embeddable": su
# archivo python3xx._pth reemplaza el cálculo normal de sys.path y, con él,
# se pierde el añadido automático de la carpeta del script. Sin esto,
# "import fs_contrato" falla aunque el módulo esté al lado.
_AQUI = Path(getattr(sys, "_MEIPASS", Path(__file__).resolve().parent))
if str(_AQUI) not in sys.path:
    sys.path.insert(0, str(_AQUI))

import fs_contrato as C

XMLSPACE = "{http://www.w3.org/XML/1998/namespace}space"

# Anchos por defecto de las 4 columnas, en twips (1/1440"). Suman 9360 =
# 6.5", el ancho útil de una carta con márgenes de 1".
ANCHOS_DEFECTO = (4500, 900, 1980, 1980)

ESTILO_ETIQUETA = {
    "H": {"negrita": True, "sangria": 0, "espacio_antes": 120},
    "I": {"negrita": False, "sangria": 220, "espacio_antes": 0},
    "S": {"negrita": False, "sangria": 220, "espacio_antes": 0},
    "T": {"negrita": True, "sangria": 0, "espacio_antes": 60},
    "N": {"negrita": False, "sangria": 0, "espacio_antes": 120, "cursiva": True},
}


# --------------------------------------------------------------------------- #
#  Primitivas OOXML
# --------------------------------------------------------------------------- #
def _id_estable(tag):
    """Un w:id determinista por tag: dos construcciones seguidas producen el
    mismo XML (condición necesaria para que el refresco sea idempotente)."""
    h = hashlib.md5(tag.encode("utf-8")).hexdigest()[:8]
    return (int(h, 16) % 2_000_000_000) + 1000


def _run(texto_, negrita=None, cursiva=False, oculto=False):
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    usa = False
    if negrita is not None:
        b = OxmlElement("w:b")
        if not negrita:
            b.set(qn("w:val"), "0")
        rPr.append(b)
        usa = True
    if cursiva:
        rPr.append(OxmlElement("w:i"))
        usa = True
    if oculto:
        rPr.append(OxmlElement("w:vanish"))
        usa = True
    if usa:
        r.append(rPr)
    t = OxmlElement("w:t")
    t.text = "" if texto_ is None else str(texto_)
    t.set(XMLSPACE, "preserve")
    r.append(t)
    return r


def _parrafo(estilo=None, sangria=None, alineacion=None, espacio_antes=None):
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    usa = False
    if estilo:
        ps = OxmlElement("w:pStyle")
        ps.set(qn("w:val"), estilo)
        pPr.append(ps)
        usa = True
    if espacio_antes:
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:before"), str(espacio_antes))
        pPr.append(sp)
        usa = True
    if sangria:
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(sangria))
        pPr.append(ind)
        usa = True
    if alineacion:
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), alineacion)
        pPr.append(jc)
        usa = True
    if usa:
        p.append(pPr)
    return p


def _parrafo_texto(texto_, **kw):
    negrita = kw.pop("negrita", None)
    cursiva = kw.pop("cursiva", False)
    oculto = kw.pop("oculto", False)
    p = _parrafo(**kw)
    p.append(_run(texto_, negrita=negrita, cursiva=cursiva, oculto=oculto))
    return p


def _borde(nombre, tipo="single", sz="4"):
    b = OxmlElement(nombre)
    b.set(qn("w:val"), tipo)
    b.set(qn("w:sz"), sz)
    b.set(qn("w:space"), "0")
    b.set(qn("w:color"), "auto")
    return b


def _celda(ancho, parrafos, borde_sup=None, borde_inf=None):
    tc = OxmlElement("w:tc")
    tcPr = OxmlElement("w:tcPr")
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(ancho))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)
    if borde_sup or borde_inf:
        bs = OxmlElement("w:tcBorders")
        if borde_sup:
            bs.append(_borde("w:top", borde_sup))
        if borde_inf:
            bs.append(_borde("w:bottom", borde_inf))
        tcPr.append(bs)
    tc.append(tcPr)
    for p in parrafos:
        tc.append(p)
    return tc


def _sdt(tag, alias=None, bloqueado=True, en_linea=False):
    """Crea un control de contenido vacío con su etiqueta (Tag).

    bloqueado -> w:lock="sdtContentLocked": Word impide editar el interior a
    mano. No estorba a este motor (escribimos el XML directamente), pero sí
    obliga al add-in a desbloquear antes de escribir.
    """
    sdt = OxmlElement("w:sdt")
    pr = OxmlElement("w:sdtPr")
    if alias:
        a = OxmlElement("w:alias")
        a.set(qn("w:val"), alias)
        pr.append(a)
    t = OxmlElement("w:tag")
    t.set(qn("w:val"), tag)
    pr.append(t)
    i = OxmlElement("w:id")
    i.set(qn("w:val"), str(_id_estable(tag)))
    pr.append(i)
    if bloqueado:
        lk = OxmlElement("w:lock")
        lk.set(qn("w:val"), "sdtContentLocked")
        pr.append(lk)
    pr.append(OxmlElement("w:text") if en_linea else OxmlElement("w:richText"))
    sdt.append(pr)
    sdt.append(OxmlElement("w:sdtContent"))
    return sdt


def _contenido(sdt):
    return sdt.find(qn("w:sdtContent"))


def _tag_de(sdt):
    pr = sdt.find(qn("w:sdtPr"))
    if pr is None:
        return None
    t = pr.find(qn("w:tag"))
    return t.get(qn("w:val")) if t is not None else None


def _indexar(doc):
    """{tag: elemento sdt} de todo el documento. Ante duplicados, el primero."""
    idx = {}
    for sdt in doc.element.body.iter(qn("w:sdt")):
        tag = _tag_de(sdt)
        if tag and tag not in idx:
            idx[tag] = sdt
    return idx


def _cuerpo_append(doc, el):
    """Añade al final del cuerpo, siempre ANTES del w:sectPr final."""
    body = doc.element.body
    sect = body.find(qn("w:sectPr"))
    if sect is not None:
        sect.addprevious(el)
    else:
        body.append(el)


def _vaciar(el):
    for hijo in list(el):
        el.remove(hijo)


# --------------------------------------------------------------------------- #
#  Tablas
# --------------------------------------------------------------------------- #
def _anchos_de(tbl):
    """Lee los anchos del w:tblGrid. Si el usuario arrastra las columnas en
    Word, el tblGrid se actualiza y el refresco respeta la nueva medida."""
    grid = tbl.find(qn("w:tblGrid")) if tbl is not None else None
    if grid is None:
        return list(ANCHOS_DEFECTO)
    anchos = []
    for gc in grid.findall(qn("w:gridCol")):
        try:
            anchos.append(int(gc.get(qn("w:w"))))
        except (TypeError, ValueError):
            anchos.append(ANCHOS_DEFECTO[len(anchos) % 4])
    return anchos or list(ANCHOS_DEFECTO)


def _tblPr_defecto():
    tblPr = OxmlElement("w:tblPr")
    w = OxmlElement("w:tblW")
    w.set(qn("w:w"), "0")
    w.set(qn("w:type"), "auto")
    tblPr.append(w)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    return tblPr


def _tblGrid(anchos):
    grid = OxmlElement("w:tblGrid")
    for a in anchos:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(a))
        grid.append(gc)
    return grid


def _fila_encabezado(ctx, anchos):
    tr = OxmlElement("w:tr")
    trPr = OxmlElement("w:trPr")
    trPr.append(OxmlElement("w:tblHeader"))
    tr.append(trPr)

    fa = str(ctx.get("fecha_actual", "") or "")
    fp = str(ctx.get("fecha_previa", "") or "")
    ea = str(ctx.get("estado_actual", "") or "")
    ep = str(ctx.get("estado_previo", "") or "")

    def cab(texto_, sub, alineacion):
        ps = [_parrafo_texto(texto_, negrita=True, alineacion=alineacion)]
        if sub:
            ps.append(_parrafo_texto(sub, negrita=False, cursiva=True, alineacion=alineacion))
        return ps

    tr.append(_celda(anchos[0], cab("", "", None), borde_inf="single"))
    tr.append(_celda(anchos[1], cab("Nota", "", "center"), borde_inf="single"))
    tr.append(_celda(anchos[2], cab(fa, ea, "right"), borde_inf="single"))
    tr.append(_celda(anchos[3], cab(fp, ep, "right"), borde_inf="single"))
    return tr


def _fila_de_linea(linea, anchos):
    tipo = linea.get("tipo", "I")
    est = ESTILO_ETIQUETA.get(tipo, ESTILO_ETIQUETA["I"])
    tr = OxmlElement("w:tr")

    borde_sup = "single" if tipo in ("S", "T") else None
    borde_inf = "double" if tipo == "T" else None

    etiqueta = str(linea.get("etiqueta", "") or "")
    nota = str(linea.get("nota", "") or "")
    actual = str(linea.get("actual", "") or "")
    previo = str(linea.get("previo", "") or "")

    if tipo in ("H", "N"):
        nota = actual = previo = ""

    p_et = _parrafo_texto(
        etiqueta,
        negrita=est["negrita"],
        cursiva=est.get("cursiva", False),
        sangria=est["sangria"],
        espacio_antes=est["espacio_antes"],
    )
    tr.append(_celda(anchos[0], [p_et], borde_sup=borde_sup, borde_inf=borde_inf))
    tr.append(
        _celda(
            anchos[1],
            [_parrafo_texto(nota, alineacion="center", espacio_antes=est["espacio_antes"])],
            borde_sup=borde_sup,
            borde_inf=borde_inf,
        )
    )
    for ancho, val in ((anchos[2], actual), (anchos[3], previo)):
        tr.append(
            _celda(
                ancho,
                [
                    _parrafo_texto(
                        val,
                        negrita=est["negrita"],
                        alineacion="right",
                        espacio_antes=est["espacio_antes"],
                    )
                ],
                borde_sup=borde_sup,
                borde_inf=borde_inf,
            )
        )
    return tr


def _tabla(ctx, lineas, anchos=None, tblPr=None):
    anchos = list(anchos or ANCHOS_DEFECTO)
    while len(anchos) < 4:
        anchos.append(ANCHOS_DEFECTO[len(anchos)])
    tbl = OxmlElement("w:tbl")
    tbl.append(copy.deepcopy(tblPr) if tblPr is not None else _tblPr_defecto())
    tbl.append(_tblGrid(anchos[:4]))
    tbl.append(_fila_encabezado(ctx, anchos))
    for linea in lineas:
        tbl.append(_fila_de_linea(linea, anchos))
    return tbl


def _lineas_de_tabla(nombre, ctx):
    """Qué líneas del Excel alimentan la tabla `fs-tabla-<nombre>`.

    'principal' -> todas.
    Cualquier otro nombre -> las líneas de la sección cuyo encabezado (fila
    de tipo H) produce esa misma clave. Permite partir el estado en varias
    tablas sin tocar el código: basta con nombrar el control de contenido
    'fs-tabla-current_assets' y esa tabla recibe solo esa sección.
    """
    lineas = ctx.get("lineas", [])
    if nombre == C.TABLA_PRINCIPAL:
        return list(lineas)

    seleccion = []
    dentro = False
    for linea in lineas:
        if linea.get("tipo") == "H":
            dentro = C.clave(linea.get("etiqueta")) == nombre
            if dentro:
                seleccion.append(linea)
            continue
        if dentro:
            seleccion.append(linea)
    return seleccion


# --------------------------------------------------------------------------- #
#  Construcción / reparación del andamiaje
# --------------------------------------------------------------------------- #
def _bloque_prosa(nombre, texto_guia):
    sdt = _sdt(
        C.tag_prosa(nombre),
        alias=f"Redacción — {nombre}",
        bloqueado=False,          # ESTA es la zona que la persona edita
    )
    cont = _contenido(sdt)
    cont.append(_parrafo_texto(texto_guia))
    # Un párrafo vacío de holgura. Con el documento protegido, el rango
    # editable termina en la última marca de párrafo: sin este hueco, quien
    # solo tiene el rol Redactor puede corregir el texto existente pero no
    # empezar un párrafo nuevo al final de la zona.
    cont.append(_parrafo())
    return sdt


def _bloque_campo(nombre, valor=""):
    sdt = _sdt(C.tag_campo(nombre), alias=f"Campo — {nombre}", en_linea=True)
    _contenido(sdt).append(_run(valor or f"«{nombre}»"))
    return sdt


def construir(ruta, ctx=None, verbose=True):
    """Añade al documento lo que le falte para cumplir el contrato.

    Idempotente: correrla dos veces no duplica nada. No borra prosa ni
    reordena lo que ya exista; solo agrega las anclas ausentes al final.
    """
    doc = Document(str(ruta))
    idx = _indexar(doc)
    añadidos = []
    ctx = ctx or {}

    def falta(tag):
        return tag not in idx

    # --- encabezado del estado -------------------------------------------- #
    #  Cada campo va dentro de una frase con sentido, no en una línea suelta:
    #  el documento tiene que poder imprimirse tal cual desde el primer día.
    #  Un campo que ya exista en el documento no se vuelve a poner, así que
    #  la persona puede moverlos donde quiera y 'reparar' los respeta.
    def linea_campos(piezas, **kw):
        """piezas: lista de str (texto fijo) o ('campo', nombre)."""
        nombres = [x[1] for x in piezas if isinstance(x, tuple)]
        if not any(falta(C.tag_campo(n)) for n in nombres):
            return                      # todos sus campos ya existen: no repetir
        p = _parrafo(**kw)
        for pieza in piezas:
            if isinstance(pieza, tuple):
                nombre = pieza[1]
                if falta(C.tag_campo(nombre)):
                    p.append(_bloque_campo(nombre, str(ctx.get(nombre, ""))))
                    añadidos.append(C.tag_campo(nombre))
                else:
                    p.append(_run(str(ctx.get(nombre, "")) or f"«{nombre}»"))
            else:
                p.append(_run(pieza))
        _cuerpo_append(doc, p)

    linea_campos([("campo", "empresa")], alineacion="center", espacio_antes=240)
    linea_campos([("campo", "titulo")], alineacion="center")
    linea_campos(
        ["Al ", ("campo", "fecha_actual"), " (", ("campo", "estado_actual"),
         ") — comparado con ", ("campo", "fecha_previa"),
         " (", ("campo", "estado_previo"), ")"],
        alineacion="center",
    )
    linea_campos(
        ["Cifras expresadas en ", ("campo", "moneda"), ", en unidades de ",
         ("campo", "miles")],
        alineacion="center",
    )

    # --- zonas de prosa + tabla ------------------------------------------- #
    if falta(C.tag_prosa("introduccion")):
        _cuerpo_append(
            doc,
            _bloque_prosa(
                "introduccion",
                "Zona de redacción libre. Escriba aquí lo que quiera: este "
                "texto NO se toca al refrescar las cifras.",
            ),
        )
        añadidos.append(C.tag_prosa("introduccion"))

    tag_tabla = C.tag_tabla(C.TABLA_PRINCIPAL)
    if falta(tag_tabla):
        sdt = _sdt(tag_tabla, alias="Tabla — estado principal")
        lineas = _lineas_de_tabla(C.TABLA_PRINCIPAL, ctx) if ctx else []
        _contenido(sdt).append(_tabla(ctx, lineas))
        _cuerpo_append(doc, sdt)
        # Word necesita un párrafo tras una tabla al final del cuerpo
        _cuerpo_append(doc, _parrafo())
        añadidos.append(tag_tabla)

    if falta(C.tag_prosa("analisis")):
        _cuerpo_append(
            doc,
            _bloque_prosa(
                "analisis",
                "Zona de análisis. Aquí puede intercalar cifras vivas: use "
                "«fs_documento.py catalogo» para ver las claves disponibles.",
            ),
        )
        añadidos.append(C.tag_prosa("analisis"))

    # --- bitácora y foto -------------------------------------------------- #
    if falta(C.TAG_REGISTRO):
        sdt = _sdt(C.TAG_REGISTRO, alias="Bitácora de actualizaciones", bloqueado=True)
        _contenido(sdt).append(_parrafo_texto("Bitácora de actualizaciones", negrita=True))
        _cuerpo_append(doc, sdt)
        añadidos.append(C.TAG_REGISTRO)

    if falta(C.TAG_META):
        sdt = _sdt(C.TAG_META, alias="Metadatos (oculto)", bloqueado=True)
        _contenido(sdt).append(_parrafo_texto("{}", oculto=True))
        _cuerpo_append(doc, sdt)
        añadidos.append(C.TAG_META)

    doc.save(str(ruta))
    if verbose:
        if añadidos:
            print(f"  Anclas añadidas ({len(añadidos)}):")
            for t in añadidos:
                print(f"    + {t}")
        else:
            print("  Nada que añadir: el documento ya cumple el contrato.")
    return añadidos


# --------------------------------------------------------------------------- #
#  Refresco en el sitio
# --------------------------------------------------------------------------- #
def _escribir_en_linea(sdt, valor):
    """Reemplaza el texto de un control en línea conservando su formato."""
    cont = _contenido(sdt)
    rPr = None
    for r in cont.iter(qn("w:r")):
        hallado = r.find(qn("w:rPr"))
        if hallado is not None:
            rPr = copy.deepcopy(hallado)
        break
    _vaciar(cont)
    r = OxmlElement("w:r")
    if rPr is not None:
        r.append(rPr)
    t = OxmlElement("w:t")
    t.text = "" if valor is None else str(valor)
    t.set(XMLSPACE, "preserve")
    r.append(t)
    cont.append(r)


def _leer_en_linea(sdt):
    return "".join(t.text or "" for t in _contenido(sdt).iter(qn("w:t")))


def _leer_meta(idx):
    sdt = idx.get(C.TAG_META)
    if sdt is None:
        return {}
    crudo = "".join(t.text or "" for t in _contenido(sdt).iter(qn("w:t")))
    try:
        return json.loads(crudo) or {}
    except (json.JSONDecodeError, TypeError):
        return {}


def _guardar_meta(idx, ctx, origen):
    sdt = idx.get(C.TAG_META)
    if sdt is None:
        return
    foto = {
        "fecha": datetime.now().isoformat(timespec="seconds"),
        "origen": origen,
        "lineas": [
            {
                "e": l.get("etiqueta", ""),
                "t": l.get("tipo", ""),
                "a": l.get("actual", ""),
                "p": l.get("previo", ""),
            }
            for l in ctx.get("lineas", [])
        ],
    }
    cont = _contenido(sdt)
    _vaciar(cont)
    cont.append(
        _parrafo_texto(json.dumps(foto, ensure_ascii=False, separators=(",", ":")), oculto=True)
    )


def _claves_diff(filas, etiqueta_de, tipo_de):
    """Un nombre estable y legible para cada fila, apto para comparar.

    Las filas de subtotal (S) no traen etiqueta: todas se llamarían igual y
    el diff las daría por nuevas en cada corrida. Se las nombra por la
    sección en la que caen ('Subtotal de Current assets:'), y si aun así
    dos coinciden se numeran.
    """
    nombres = []
    seccion = ""
    vistos = {}
    for f in filas:
        etiqueta = (etiqueta_de(f) or "").strip()
        tipo = tipo_de(f) or ""
        if tipo == "H" and etiqueta:
            seccion = etiqueta.rstrip(":").strip()
        if etiqueta:
            base = etiqueta
        elif tipo == "S":
            base = f"Subtotal de {seccion}" if seccion else "Subtotal"
        else:
            base = f"(sin etiqueta {tipo})"
        vistos[base] = vistos.get(base, 0) + 1
        nombres.append(base if vistos[base] == 1 else f"{base} #{vistos[base]}")
    return nombres


def _calcular_cambios(meta_previa, ctx):
    previas = meta_previa.get("lineas") or []
    if not previas:
        return ["Primera actualización: no hay versión anterior con la que comparar."]

    nom_antes = _claves_diff(previas, lambda d: d.get("e"), lambda d: d.get("t"))
    nom_ahora = _claves_diff(
        ctx.get("lineas", []), lambda d: d.get("etiqueta"), lambda d: d.get("tipo")
    )
    antes = {n: f for n, f in zip(nom_antes, previas)}

    cambios = []
    for nombre, l in zip(nom_ahora, ctx.get("lineas", [])):
        a = antes.pop(nombre, None)
        if a is None:
            cambios.append(f"Nueva fila: {nombre}  {l.get('actual') or '—'}")
        elif a.get("a") != l.get("actual") or a.get("p") != l.get("previo"):
            cambios.append(
                f"{nombre}: {a.get('a') or '—'} → {l.get('actual') or '—'}"
                f"  (comparativo {a.get('p') or '—'} → {l.get('previo') or '—'})"
            )
    for nombre in antes:
        cambios.append(f"Fila retirada: {nombre}")
    return cambios or ["Sin cambios en las cifras respecto de la última actualización."]


def _escribir_registro(idx, cambios, origen):
    sdt = idx.get(C.TAG_REGISTRO)
    if sdt is None:
        return False
    cont = _contenido(sdt)
    sello = datetime.now().strftime("%Y-%m-%d %H:%M")
    bloque = [
        _parrafo_texto(f"Actualización {sello} — origen: {origen}", negrita=True,
                       espacio_antes=120)
    ]
    for c in cambios[:40]:
        bloque.append(_parrafo_texto(f"• {c}", sangria=220))
    if len(cambios) > 40:
        bloque.append(_parrafo_texto(f"• … y {len(cambios) - 40} cambios más.", sangria=220))

    titulo = cont.find(qn("w:p"))
    ancla = titulo if titulo is not None else None
    for el in reversed(bloque):
        if ancla is not None:
            ancla.addnext(el)
        else:
            cont.append(el)
    return True


def insertar_dato(ruta, clave_, campo, zona="analisis", antes="", despues="",
                  valor_inicial="—", verbose=True):
    """Añade un párrafo con una cifra viva al final de una zona de prosa.

    Es el equivalente por línea de órdenes al botón «insertar dato» del
    add-in: deja el control de contenido en línea, bloqueado y con la
    etiqueta correcta, listo para que el siguiente refresco lo rellene.
    """
    if campo not in C.CAMPOS_DATO:
        raise ValueError(
            f"'{campo}' no es un campo válido.\n"
            f"Válidos: {', '.join(C.CAMPOS_DATO)}"
        )
    doc = Document(str(ruta))
    idx = _indexar(doc)
    sdt_zona = idx.get(C.tag_prosa(zona))
    if sdt_zona is None:
        raise ValueError(
            f"No existe la zona de prosa 'fs-prosa-{zona}' en el documento.\n"
            f"Zonas disponibles: "
            f"{', '.join(sorted(n for t in idx for f, n, _ in [C.descomponer(t)] if f == C.FAM_PROSA)) or '(ninguna)'}"
        )

    tag = C.tag_dato(clave_, campo)
    if tag in idx:
        if verbose:
            print(f"  El ancla {tag} ya existe en el documento; no se duplica.")
        return False

    p = _parrafo()
    if antes:
        p.append(_run(antes))
    sdt = _sdt(tag, alias=f"{clave_} ({campo})", en_linea=True)
    _contenido(sdt).append(_run(valor_inicial))
    p.append(sdt)
    if despues:
        p.append(_run(despues))
    _contenido(sdt_zona).append(p)
    doc.save(str(ruta))
    if verbose:
        print(f"  + {tag}  en la zona '{zona}'")
    return True


def refrescar(ruta, ctx, origen="", con_registro=True, verbose=True):
    """Reescribe SOLO las regiones de datos. Devuelve un informe."""
    doc = Document(str(ruta))
    idx = _indexar(doc)
    valores, colisiones = C.construir_valores(ctx)

    meta_previa = _leer_meta(idx)
    cambios = _calcular_cambios(meta_previa, ctx)

    informe = {
        "tablas": [], "campos": 0, "datos": 0,
        "huerfanos": [], "colisiones": colisiones,
        "cambios": cambios, "sin_ancla_prosa": 0,
    }

    for tag, sdt in idx.items():
        familia, nombre, campo = C.descomponer(tag)

        if familia == C.FAM_TABLA:
            cont = _contenido(sdt)
            vieja = cont.find(qn("w:tbl"))
            anchos = _anchos_de(vieja)
            tblPr = vieja.find(qn("w:tblPr")) if vieja is not None else None
            lineas = _lineas_de_tabla(nombre, ctx)
            nueva = _tabla(ctx, lineas, anchos=anchos, tblPr=tblPr)
            _vaciar(cont)
            cont.append(nueva)
            informe["tablas"].append((nombre, len(lineas)))

        elif familia == C.FAM_CAMPO:
            if tag in valores:
                _escribir_en_linea(sdt, valores[tag])
                informe["campos"] += 1
            else:
                informe["huerfanos"].append(tag)

        elif familia == C.FAM_DATO:
            if tag in valores:
                _escribir_en_linea(sdt, valores[tag])
                informe["datos"] += 1
            else:
                informe["huerfanos"].append(tag)

        elif familia == C.FAM_PROSA:
            informe["sin_ancla_prosa"] += 1

    if con_registro:
        informe["con_registro"] = _escribir_registro(idx, cambios, origen)
    else:
        informe["con_registro"] = False

    _guardar_meta(idx, ctx, origen)
    doc.save(str(ruta))
    return informe


# --------------------------------------------------------------------------- #
#  Verificación
# --------------------------------------------------------------------------- #
def verificar(ruta, ctx=None):
    doc = Document(str(ruta))
    idx = _indexar(doc)
    valores = C.construir_valores(ctx)[0] if ctx else {}

    rep = {
        "tablas": [], "campos": [], "datos": [], "prosa": [],
        "registro": C.TAG_REGISTRO in idx, "meta": C.TAG_META in idx,
        "desconocidos": [], "huerfanos": [], "sin_usar": [],
    }

    for tag, sdt in idx.items():
        familia, nombre, campo = C.descomponer(tag)
        if familia == C.FAM_TABLA:
            cont = _contenido(sdt)
            tiene = cont.find(qn("w:tbl")) is not None
            rep["tablas"].append((nombre, tiene))
        elif familia == C.FAM_CAMPO:
            rep["campos"].append((nombre, _leer_en_linea(sdt)))
            if valores and tag not in valores:
                rep["huerfanos"].append(tag)
        elif familia == C.FAM_DATO:
            rep["datos"].append((nombre, campo, _leer_en_linea(sdt)))
            if valores and tag not in valores:
                rep["huerfanos"].append(tag)
        elif familia == C.FAM_PROSA:
            rep["prosa"].append(nombre)
        elif familia in ("registro", "meta"):
            pass
        else:
            rep["desconocidos"].append(tag)

    if valores:
        usados = set(idx)
        rep["sin_usar"] = sorted(
            t for t in valores
            if t.startswith("fs-dato-") and t not in usados and valores[t]
        )
    return rep


# --------------------------------------------------------------------------- #
#  Protección (los dos editores)
# --------------------------------------------------------------------------- #
#  Orden de w:settings según ECMA-376; documentProtection debe ir en su sitio
#  o Word se queja al abrir.
_ORDEN_SETTINGS = [
    "writeProtection", "view", "zoom", "removePersonalInformation",
    "removeDateAndTime", "doNotDisplayPageBoundaries", "displayBackgroundShape",
    "printPostScriptOverText", "printFractionalCharacterWidth", "printFormsData",
    "embedTrueTypeFonts", "embedSystemFonts", "saveSubsetFonts", "saveFormsData",
    "mirrorMargins", "alignBordersAndEdges", "bordersDoNotSurroundHeader",
    "bordersDoNotSurroundFooter", "gutterAtTop", "hideSpellingErrors",
    "hideGrammaticalErrors", "activeWritingStyle", "proofState", "formsDesign",
    "attachedTemplate", "linkStyles", "stylePaneFormatFilter",
    "stylePaneSortMethod", "documentType", "mailMerge", "revisionView",
    "trackChanges", "doNotTrackMoves", "doNotTrackFormatting",
    "documentProtection",
]


def _hash_proteccion(clave_, salt, vueltas=100000):
    """Algoritmo de ECMA-376 para w:hash (SHA-1, cryptAlgorithmSid=4)."""
    h = hashlib.sha1(salt + clave_.encode("utf-16-le")).digest()
    for i in range(vueltas):
        h = hashlib.sha1(h + struct.pack("<I", i)).digest()
    return h


def proteger(ruta, clave_, verbose=True):
    """Deja el documento en solo lectura salvo las zonas fs-prosa-*.

    Rol REDACTOR  : abre y solo puede escribir dentro de las zonas de prosa.
    Rol EDITOR DE DATOS : conoce la clave (o usa el add-in), que desprotege,
                          refresca y vuelve a proteger.
    """
    doc = Document(str(ruta))

    # 1. rangos editables alrededor de cada zona de prosa
    n = 0
    for sdt in doc.element.body.iter(qn("w:sdt")):
        tag = _tag_de(sdt)
        familia, _, _ = C.descomponer(tag)
        if familia != C.FAM_PROSA:
            continue
        cont = _contenido(sdt)
        if cont.find(qn("w:permStart")) is not None:
            continue
        n += 1
        ini = OxmlElement("w:permStart")
        ini.set(qn("w:id"), str(n))
        ini.set(qn("w:edGrp"), "everyone")
        fin = OxmlElement("w:permEnd")
        fin.set(qn("w:id"), str(n))
        cont.insert(0, ini)
        cont.append(fin)

    # 2. protección global
    settings = doc.settings.element
    for viejo in settings.findall(qn("w:documentProtection")):
        settings.remove(viejo)

    salt = os.urandom(16)
    prot = OxmlElement("w:documentProtection")
    prot.set(qn("w:edit"), "readOnly")
    prot.set(qn("w:enforcement"), "1")
    prot.set(qn("w:cryptProviderType"), "rsaFull")
    prot.set(qn("w:cryptAlgorithmClass"), "hash")
    prot.set(qn("w:cryptAlgorithmType"), "typeAny")
    prot.set(qn("w:cryptAlgorithmSid"), "4")
    prot.set(qn("w:cryptSpinCount"), "100000")
    prot.set(qn("w:hash"), base64.b64encode(_hash_proteccion(clave_, salt)).decode())
    prot.set(qn("w:salt"), base64.b64encode(salt).decode())

    pos = _ORDEN_SETTINGS.index("documentProtection")
    posteriores = set(_ORDEN_SETTINGS[pos + 1:])
    ancla = None
    for hijo in settings:
        nombre = hijo.tag.split("}")[-1]
        if nombre in posteriores or nombre not in _ORDEN_SETTINGS:
            ancla = hijo
            break
    if ancla is not None:
        ancla.addprevious(prot)
    else:
        settings.append(prot)

    doc.save(str(ruta))
    if verbose:
        print(f"  Protegido. Zonas de prosa editables: {n}")
    return n


def desproteger(ruta, verbose=True):
    doc = Document(str(ruta))
    settings = doc.settings.element
    quitados = 0
    for viejo in settings.findall(qn("w:documentProtection")):
        settings.remove(viejo)
        quitados += 1
    doc.save(str(ruta))
    if verbose:
        print("  Protección retirada." if quitados else "  No estaba protegido.")
    return quitados


# --------------------------------------------------------------------------- #
#  Utilidades de línea de órdenes
# --------------------------------------------------------------------------- #
def _cargar_ctx(argumento=None):
    """Lee el Excel con el motor de generador_fs y devuelve su contexto."""
    import generador_fs as G

    cfg = G.cargar_config()
    xlsx = Path(argumento).resolve() if argumento else G.encontrar_excel_por_convencion(cfg)
    if not xlsx.exists():
        raise ValueError(f"No se encontró el libro de Excel:\n  {xlsx}")
    ctx = G.leer_contexto(xlsx, cfg)
    ctx.pop("_meta", None)
    ctx.pop("_avisos", None)
    return ctx, xlsx


def _respaldar(ruta):
    ruta = Path(ruta)
    bak = ruta.with_suffix(ruta.suffix + ".bak")
    shutil.copy2(ruta, bak)
    return bak


def _titulo(t):
    print()
    print("=" * 68)
    print(f" {t}")
    print("=" * 68)


def main(argv):
    if len(argv) < 2:
        print(__doc__)
        return 1

    orden = argv[1].lower()
    args = [a for a in argv[2:] if not a.startswith("--")]
    flags = {a.lower() for a in argv[2:] if a.startswith("--")}

    def opcion(nombre, defecto=None):
        for i, a in enumerate(argv):
            if a.lower() == nombre and i + 1 < len(argv):
                return argv[i + 1]
        return defecto

    if orden == "catalogo":
        ctx, xlsx = _cargar_ctx(args[0] if args else None)
        _titulo(f"CIFRAS DISPONIBLES — {xlsx.name}")
        print(f" {'clave':42} {'actual':>16} {'previo':>16}")
        print(" " + "-" * 76)
        for k, etiqueta, actual, previo in C.catalogo(ctx):
            print(f" {k:42} {actual:>16} {previo:>16}")
        print()
        print(" Para intercalar una cifra en la prosa, inserte en Word un control")
        print(" de contenido de TEXTO con la etiqueta (Tag):")
        print("     fs-dato-<clave>-actual     (o -previo, -nota, -var_abs, -var_pct)")
        print("=" * 68)
        return 0

    if orden == "plantilla":
        if not args:
            print("Falta el destino: plantilla <destino.docx>")
            return 1
        destino = Path(args[0]).resolve()
        destino.parent.mkdir(parents=True, exist_ok=True)
        Document().save(str(destino))
        _titulo(f"PLANTILLA BASE — {destino.name}")
        ctx = None
        try:
            ctx, _ = _cargar_ctx(opcion("--excel"))
        except Exception:
            pass
        construir(destino, ctx or {})
        print(f" Escrita en: {destino}")
        print("=" * 68)
        return 0

    if not args:
        print(f"Falta la ruta del documento: {orden} <doc.docx>")
        return 1
    doc_ruta = Path(args[0]).resolve()
    if not doc_ruta.exists():
        print(f"No se encontró el documento:\n  {doc_ruta}")
        return 1

    if orden in ("construir", "reparar"):
        ctx = None
        try:
            ctx, _ = _cargar_ctx(args[1] if len(args) > 1 else None)
        except Exception:
            pass
        bak = _respaldar(doc_ruta)
        _titulo(f"ANDAMIAJE — {doc_ruta.name}")
        construir(doc_ruta, ctx or {})
        print(f" Copia previa: {bak.name}")
        print("=" * 68)
        return 0

    if orden == "refrescar":
        ctx, xlsx = _cargar_ctx(args[1] if len(args) > 1 else None)
        bak = _respaldar(doc_ruta)
        sha = hashlib.sha256(xlsx.read_bytes()).hexdigest()[:12]
        inf = refrescar(
            doc_ruta, ctx, origen=f"{xlsx.name} (sha {sha})",
            con_registro="--sin-registro" not in flags,
        )
        _titulo(f"DOCUMENTO ACTUALIZADO — {doc_ruta.name}")
        print(f" Origen:        {xlsx.name}  sha256 {sha}")
        for nombre, n in inf["tablas"]:
            print(f" Tabla '{nombre}': {n} filas escritas")
        print(f" Campos de encabezado: {inf['campos']}")
        print(f" Cifras en prosa:      {inf['datos']}")
        print(f" Zonas de prosa intactas: {inf['sin_ancla_prosa']}")
        print(f" Bitácora: {'sí' if inf['con_registro'] else 'no (falta fs-registro)'}")
        print(f" Copia previa: {bak.name}")
        if inf["huerfanos"]:
            print()
            print(" AVISO — anclas en el documento sin cifra en el Excel:")
            for t in inf["huerfanos"]:
                print(f"   ? {t}")
        if inf["colisiones"]:
            print()
            print(" AVISO — etiquetas distintas que producen la misma clave:")
            for nueva, primera, k in inf["colisiones"]:
                print(f"   '{nueva}' choca con '{primera}' (clave {k}); se usó la primera.")
        print()
        print(" Cambios respecto de la última actualización:")
        for c in inf["cambios"][:25]:
            print(f"   • {c}")
        if len(inf["cambios"]) > 25:
            print(f"   • … y {len(inf['cambios']) - 25} más (ver la bitácora del documento).")
        print("=" * 68)
        return 0

    if orden == "insertar":
        if len(args) < 3:
            print("Uso: insertar <doc.docx> <clave> <campo> "
                  "[--zona <nombre>] [--antes \"texto\"] [--despues \"texto\"]")
            print(f"     campos válidos: {', '.join(C.CAMPOS_DATO)}")
            print("     para ver las claves: python fs_documento.py catalogo")
            return 1
        bak = _respaldar(doc_ruta)
        _titulo(f"INSERTAR CIFRA — {doc_ruta.name}")
        insertar_dato(
            doc_ruta, args[1], args[2],
            zona=opcion("--zona", "analisis"),
            antes=opcion("--antes", ""),
            despues=opcion("--despues", ""),
        )
        print(f" Copia previa: {bak.name}")
        print(" Corra 'refrescar' para que tome el valor del Excel.")
        print("=" * 68)
        return 0

    if orden == "verificar":
        ctx = None
        try:
            ctx, _ = _cargar_ctx(args[1] if len(args) > 1 else None)
        except Exception:
            pass
        rep = verificar(doc_ruta, ctx)
        _titulo(f"VERIFICACIÓN — {doc_ruta.name}")
        for nombre, tiene in rep["tablas"]:
            print(f" Tabla '{nombre}': {'con tabla dentro' if tiene else 'VACÍA (falta la tabla)'}")
        if not rep["tablas"]:
            print(" Tablas: NINGUNA. Corra 'construir'.")
        print(f" Campos de encabezado: {len(rep['campos'])}")
        for nombre, val in sorted(rep["campos"]):
            print(f"   {nombre:16} = {val!r}")
        print(f" Cifras en prosa: {len(rep['datos'])}")
        for nombre, campo, val in sorted(rep["datos"]):
            print(f"   {nombre}-{campo} = {val!r}")
        print(f" Zonas de prosa: {', '.join(sorted(rep['prosa'])) or '(ninguna)'}")
        print(f" Bitácora: {'sí' if rep['registro'] else 'NO'}    "
              f"Metadatos: {'sí' if rep['meta'] else 'NO'}")
        if rep["desconocidos"]:
            print(" Controles ajenos al contrato (se ignoran al refrescar):")
            for t in rep["desconocidos"]:
                print(f"   - {t}")
        if rep["huerfanos"]:
            print(" HUÉRFANOS — el documento los pide pero el Excel no los tiene:")
            for t in rep["huerfanos"]:
                print(f"   ? {t}")
        if rep["sin_usar"]:
            print(f" Cifras disponibles sin usar en la prosa: {len(rep['sin_usar'])}"
                  f"  (vea 'catalogo')")
        print("=" * 68)
        return 0

    if orden == "proteger":
        clave_ = opcion("--clave")
        if not clave_:
            print("Falta --clave <clave>. Es la que pedirá Word para desproteger.")
            return 1
        bak = _respaldar(doc_ruta)
        _titulo(f"PROTECCIÓN — {doc_ruta.name}")
        proteger(doc_ruta, clave_)
        print(f" Copia previa: {bak.name}")
        print(" Rol Redactor: solo puede escribir dentro de las zonas fs-prosa-*.")
        print("=" * 68)
        return 0

    if orden == "desproteger":
        bak = _respaldar(doc_ruta)
        _titulo(f"PROTECCIÓN — {doc_ruta.name}")
        desproteger(doc_ruta)
        print(f" Copia previa: {bak.name}")
        print("=" * 68)
        return 0

    print(f"Orden desconocida: {orden}")
    print(__doc__)
    return 1


if __name__ == "__main__":
    try:
        sys.exit(main(sys.argv))
    except ValueError as e:
        _titulo("NO SE PUDO COMPLETAR")
        print(str(e))
        print("=" * 68)
        sys.exit(1)
    except Exception:
        _titulo("ERROR INESPERADO — copie este texto para soporte")
        traceback.print_exc()
        print("=" * 68)
        sys.exit(1)
