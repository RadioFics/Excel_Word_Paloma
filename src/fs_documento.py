"""
fs_documento.py — Motor de regiones sobre un documento de Word vivo.

A diferencia de generador_fs.py (que RENDERIZA una plantilla y produce un
.docx nuevo y desechable), este módulo ACTUALIZA EN EL SITIO un documento
que ya existe: localiza las anclas del contrato (fs_contrato.py), reescribe
únicamente su interior y no visita nada más. La prosa que una persona
redacte alrededor sobrevive intacta a cada refresco; si la borra, se queda
borrada, porque el motor nunca la vuelve a inyectar.

Órdenes
-------
    python fs_documento.py construir  <doc.docx> [--desde base.docx]
    python fs_documento.py reparar    <doc.docx>
    python fs_documento.py refrescar  <doc.docx> [libro.xlsx] [--sin-registro]
    python fs_documento.py insertar   <doc.docx> <clave> <campo> [--zona n]
    python fs_documento.py verificar  <doc.docx> [libro.xlsx]
    python fs_documento.py catalogo   [libro.xlsx]
    python fs_documento.py nombrar    [libro.xlsx] [--aplicar]
    python fs_documento.py estado     [libro.xlsx]
    python fs_documento.py limpiar-bitacora <doc.docx>
    python fs_documento.py desbloquear <doc.docx> [clave]
    python fs_documento.py bloquear    <doc.docx> [clave]
    python fs_documento.py desvincular <doc.docx> <clave>
    python fs_documento.py simplificar <doc.docx> [--quitar-zonas]
    python fs_documento.py apariencia  <doc.docx> <visible|invisible>
    python fs_documento.py proteger    <doc.docx> --clave X [--salvo-datos]
    python fs_documento.py plantilla  <destino.docx> [--excel libro.xlsx]
    python fs_documento.py proteger   <doc.docx> --clave <clave>
    python fs_documento.py desproteger <doc.docx>

'construir' y 'reparar' son la misma operación: añaden lo que falte para
que el documento cumpla el contrato, sin duplicar ni borrar lo que ya haya.
Se pueden correr sobre un documento con meses de redacción encima. Si el
documento ya traía redacción, lo que se añade entra detrás de un salto de
página, como un apartado aparte.

'plantilla' es lo mismo pero partiendo de la nada: crea el .docx y le monta
el andamiaje dentro. Es lo que hace la opción 6 del menú.

ADVERTENCIA: el refresco escribe sobre el archivo indicado. Si vive en
OneDrive, ciérrelo en Word antes de refrescar (o Word y el motor pelearán
por el archivo). Siempre se deja una copia .bak junto al original.
"""
import sys
import json
import copy
import shutil
import hashlib
import struct
import base64
import os
import traceback
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsmap

# El Python portable de .\python\ es la distribución "embeddable": su
# archivo python3xx._pth reemplaza el cálculo normal de sys.path y, con él,
# se pierde el añadido automático de la carpeta del script. Sin esto,
# "import fs_contrato" falla aunque el módulo esté al lado.
_AQUI = Path(getattr(sys, "_MEIPASS", Path(__file__).resolve().parent))
if str(_AQUI) not in sys.path:
    sys.path.insert(0, str(_AQUI))

import fs_contrato as C

from openpyxl.utils import column_index_from_string, get_column_letter

XMLSPACE = "{http://www.w3.org/XML/1998/namespace}space"

# Anchos por defecto de las 4 columnas, en twips (1/1440"). Suman 9360 =
# 6.5", el ancho útil de una carta con márgenes de 1".
ANCHOS_DEFECTO = (4500, 900, 1980, 1980)

ESTILO_ETIQUETA = {
    "H": {"negrita": True, "sangria": 0, "espacio_antes": 120},
    "I": {"negrita": False, "sangria": 220, "espacio_antes": 0},
    "S": {"negrita": False, "sangria": 220, "espacio_antes": 0},
    "T": {"negrita": True, "sangria": 0, "espacio_antes": 60},
    "N": {"negrita": False, "sangria": 0, "espacio_antes": 120, "cursiva": True},
}


# --------------------------------------------------------------------------- #
#  Consola
# --------------------------------------------------------------------------- #
#: Sustituciones para consolas antiguas. La consola de Windows suele venir en
#: cp1252, que no sabe escribir la flecha ni la vineta: sin esto, imprimir la
#: bitacora revienta con UnicodeEncodeError justo cuando hay cambios que
#: contar. En el documento se sigue guardando el caracter bonito.
_SUSTITUTOS = {
    "→": "->", "•": "*", "—": "-", "–": "-",
    "…": "...", "«": '"', "»": '"',
    "“": '"', "”": '"', "‘": "'", "’": "'",
}


def preparar_consola():
    """Evita que un carácter raro tumbe el programa a mitad del informe."""
    for flujo in (sys.stdout, sys.stderr):
        try:
            flujo.reconfigure(errors="replace")
        except (AttributeError, ValueError, OSError):
            pass


def imprimible(texto):
    """Devuelve `texto` en algo que la consola actual sepa escribir."""
    texto = "" if texto is None else str(texto)
    enc = getattr(sys.stdout, "encoding", None) or "utf-8"
    try:
        texto.encode(enc)
        return texto
    except (UnicodeEncodeError, LookupError):
        pass
    for malo, bueno in _SUSTITUTOS.items():
        texto = texto.replace(malo, bueno)
    try:
        return texto.encode(enc, "replace").decode(enc, "replace")
    except (LookupError, UnicodeError):
        return texto.encode("ascii", "replace").decode("ascii")
# --------------------------------------------------------------------------- #
#  Primitivas OOXML
# --------------------------------------------------------------------------- #
def _id_estable(tag):
    """Un w:id determinista por tag: dos construcciones seguidas producen el
    mismo XML (condición necesaria para que el refresco sea idempotente)."""
    h = hashlib.md5(tag.encode("utf-8")).hexdigest()[:8]
    return (int(h, 16) % 2_000_000_000) + 1000


def _run(texto_, negrita=None, cursiva=False, oculto=False):
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    usa = False
    if negrita is not None:
        b = OxmlElement("w:b")
        if not negrita:
            b.set(qn("w:val"), "0")
        rPr.append(b)
        usa = True
    if cursiva:
        rPr.append(OxmlElement("w:i"))
        usa = True
    if oculto:
        rPr.append(OxmlElement("w:vanish"))
        usa = True
    if usa:
        r.append(rPr)
    t = OxmlElement("w:t")
    t.text = "" if texto_ is None else str(texto_)
    t.set(XMLSPACE, "preserve")
    r.append(t)
    return r


def _parrafo(estilo=None, sangria=None, alineacion=None, espacio_antes=None):
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    usa = False
    if estilo:
        ps = OxmlElement("w:pStyle")
        ps.set(qn("w:val"), estilo)
        pPr.append(ps)
        usa = True
    if espacio_antes:
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:before"), str(espacio_antes))
        pPr.append(sp)
        usa = True
    if sangria:
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(sangria))
        pPr.append(ind)
        usa = True
    if alineacion:
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), alineacion)
        pPr.append(jc)
        usa = True
    if usa:
        p.append(pPr)
    return p


def _parrafo_texto(texto_, **kw):
    negrita = kw.pop("negrita", None)
    cursiva = kw.pop("cursiva", False)
    oculto = kw.pop("oculto", False)
    p = _parrafo(**kw)
    p.append(_run(texto_, negrita=negrita, cursiva=cursiva, oculto=oculto))
    return p


def _borde(nombre, tipo="single", sz="4"):
    b = OxmlElement(nombre)
    b.set(qn("w:val"), tipo)
    b.set(qn("w:sz"), sz)
    b.set(qn("w:space"), "0")
    b.set(qn("w:color"), "auto")
    return b


def _celda(ancho, parrafos, borde_sup=None, borde_inf=None):
    tc = OxmlElement("w:tc")
    tcPr = OxmlElement("w:tcPr")
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(ancho))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)
    if borde_sup or borde_inf:
        bs = OxmlElement("w:tcBorders")
        if borde_sup:
            bs.append(_borde("w:top", borde_sup))
        if borde_inf:
            bs.append(_borde("w:bottom", borde_inf))
        tcPr.append(bs)
    tc.append(tcPr)
    for p in parrafos:
        tc.append(p)
    return tc


def _sdt(tag, alias=None, bloqueado=True, en_linea=False):
    """Crea un control de contenido vacío con su etiqueta (Tag).

    bloqueado -> w:lock="sdtContentLocked": Word impide editar el interior a
    mano. No estorba a este motor (escribimos el XML directamente), pero sí
    obliga al add-in a desbloquear antes de escribir.
    """
    sdt = OxmlElement("w:sdt")
    pr = OxmlElement("w:sdtPr")
    if alias:
        a = OxmlElement("w:alias")
        a.set(qn("w:val"), alias)
        pr.append(a)
    t = OxmlElement("w:tag")
    t.set(qn("w:val"), tag)
    pr.append(t)
    i = OxmlElement("w:id")
    i.set(qn("w:val"), str(_id_estable(tag)))
    pr.append(i)
    if bloqueado:
        lk = OxmlElement("w:lock")
        lk.set(qn("w:val"), "sdtContentLocked")
        pr.append(lk)
    pr.append(OxmlElement("w:text") if en_linea else OxmlElement("w:richText"))
    sdt.append(pr)
    sdt.append(OxmlElement("w:sdtContent"))
    return sdt


def _contenido(sdt):
    return sdt.find(qn("w:sdtContent"))


def _tag_de(sdt):
    pr = sdt.find(qn("w:sdtPr"))
    if pr is None:
        return None
    t = pr.find(qn("w:tag"))
    return t.get(qn("w:val")) if t is not None else None


def _indexar(doc):
    """{tag: elemento sdt} de todo el documento. Ante duplicados, el primero."""
    idx = {}
    for sdt in doc.element.body.iter(qn("w:sdt")):
        tag = _tag_de(sdt)
        if tag and tag not in idx:
            idx[tag] = sdt
    return idx


def _cuerpo_append(doc, el):
    """Añade al final del cuerpo, siempre ANTES del w:sectPr final."""
    body = doc.element.body
    sect = body.find(qn("w:sectPr"))
    if sect is not None:
        sect.addprevious(el)
    else:
        body.append(el)


def _vaciar(el):
    for hijo in list(el):
        el.remove(hijo)


# --------------------------------------------------------------------------- #
#  PowerShell: la salida SIEMPRE en UTF-8
# --------------------------------------------------------------------------- #
#: Cabecera que se antepone a todos los scripts de PowerShell.
#:
#: Windows PowerShell 5.1 no escribe su salida en Unicode: usa la pagina de
#: codigos de la consola (cp437 / cp850 en un Windows en español). Cuando esa
#: salida se redirige a una tuberia, los caracteres que la pagina NO tiene no
#: dan error: se sustituyen en silencio por el parecido mas cercano.
#:
#: Eso convertia una ruta perfectamente valida en una que no existe:
#:
#:     ...EDICIÓN<espacio duro>- PAMELA.docx    (el nombre real)
#:     ...EDICION<byte 0xFF>- PAMELA.docx       (lo que llegaba)
#:              ^ perdio la tilde   ^ y el espacio duro se rompio
#:
#: y por eso un documento con el que se llevaba trabajando semanas se
#: rechazaba con «No existe» en cuanto se elegia desde el explorador.
#:
#: Los argumentos que van HACIA PowerShell viajan por CreateProcessW, que ya
#: es Unicode: ese sentido nunca ha tenido perdida.
_PS_UTF8 = (
    "try { [Console]::OutputEncoding = "
    "New-Object System.Text.UTF8Encoding $false } catch {}\n"
)


def ejecutar_ps(script_texto, *argumentos, timeout=600, sta=True):
    """Escribe un .ps1 temporal, lo ejecuta y devuelve (stdout, stderr, codigo).

    La salida se lee como BYTES y se descodifica en UTF-8: ver _PS_UTF8. Con
    `text=True` Python la descodificaba con la pagina de codigos local
    (cp1252), que no es la que PowerShell usa para escribir.

    El .ps1 se escribe con BOM porque PowerShell 5.1 lee un archivo sin BOM
    como ANSI, y entonces los acentos del propio script llegan como basura.
    """
    import subprocess
    import tempfile

    tmp = Path(tempfile.mkdtemp(prefix="fs_ps_"))
    script = tmp / "orden.ps1"
    try:
        script.write_text(_PS_UTF8 + script_texto, encoding="utf-8-sig")
        orden = ["powershell", "-NoProfile", "-ExecutionPolicy", "Bypass"]
        if sta:
            orden.append("-STA")
        orden += ["-File", str(script)]
        orden += [str(a if a is not None else "") for a in argumentos]
        res = subprocess.run(orden, capture_output=True, timeout=timeout)
    finally:
        shutil.rmtree(tmp, ignore_errors=True)

    def _texto(crudo):
        return (crudo or b"").decode("utf-8", "replace")

    return _texto(res.stdout), _texto(res.stderr), res.returncode


# --------------------------------------------------------------------------- #
#  Escritura segura
# --------------------------------------------------------------------------- #
# --------------------------------------------------------------------------- #
#  Quien tiene el archivo abierto
# --------------------------------------------------------------------------- #
def _duenos_office(ruta):
    """El nombre que Word/Excel escriben en su archivo de propietario.

    Al abrir un documento, Office crea al lado un '~$nombre.docx' cuyo primer
    byte es la longitud del nombre de usuario y el resto ese nombre. Es lo
    que Word lee para decir 'bloqueado por Fulano'. Si el documento lo tiene
    abierto OTRA PERSONA (por OneDrive/SharePoint), este suele ser el unico
    rastro visible desde aqui.
    """
    candidatos = [ruta.parent / f"~${ruta.name}"]
    # Con nombres largos Office recorta los dos primeros caracteres.
    if len(ruta.name) > 2:
        candidatos.append(ruta.parent / f"~${ruta.name[2:]}")

    crudo = b""
    for señal in candidatos:
        try:
            if señal.exists():
                crudo = señal.read_bytes()
                break
        except OSError:
            continue
    if not crudo:
        return []

    nombres = []
    n = crudo[0]
    if 0 < n < len(crudo):
        for codec in ("cp1252", "latin-1"):
            try:
                v = crudo[1:1 + n].decode(codec).strip("\x00 ").strip()
                if v and v.isprintable():
                    nombres.append(v)
                    break
            except UnicodeDecodeError:
                continue
    if not nombres and len(crudo) > 0x36:
        try:
            v = crudo[0x36:].decode("utf-16-le", errors="ignore")
            v = v.split("\x00")[0].strip()
            if v and v.isprintable():
                nombres.append(v)
        except Exception:
            pass
    return nombres


def _procesos_que_bloquean(ruta):
    """Procesos LOCALES que tienen el archivo abierto, via Restart Manager.

    Es la misma API que usa Windows para el cartel «este archivo esta siendo
    utilizado por...». No requiere permisos de administrador. Si algo falla
    (no es Windows, la DLL no esta) devuelve lista vacia en vez de romper:
    saber quien bloquea es un extra, no la comprobacion en si.
    """
    if os.name != "nt":
        return []
    try:
        import ctypes
        from ctypes import wintypes
    except ImportError:
        return []
    try:
        rstrtmgr = ctypes.WinDLL("rstrtmgr")
    except OSError:
        return []

    CCH_RM_SESSION_KEY = 32
    CCH_RM_MAX_APP_NAME = 255
    CCH_RM_MAX_SVC_NAME = 63
    ERROR_MORE_DATA = 234

    class FILETIME(ctypes.Structure):
        _fields_ = [("dwLowDateTime", wintypes.DWORD),
                    ("dwHighDateTime", wintypes.DWORD)]

    class RM_UNIQUE_PROCESS(ctypes.Structure):
        _fields_ = [("dwProcessId", wintypes.DWORD),
                    ("ProcessStartTime", FILETIME)]

    class RM_PROCESS_INFO(ctypes.Structure):
        _fields_ = [
            ("Process", RM_UNIQUE_PROCESS),
            ("strAppName", wintypes.WCHAR * (CCH_RM_MAX_APP_NAME + 1)),
            ("strServiceShortName", wintypes.WCHAR * (CCH_RM_MAX_SVC_NAME + 1)),
            ("ApplicationType", ctypes.c_uint),
            ("AppStatus", ctypes.c_ulong),
            ("TSSessionId", wintypes.DWORD),
            ("bRestartable", wintypes.BOOL),
        ]

    sesion = wintypes.DWORD(0)
    clave = ctypes.create_unicode_buffer(CCH_RM_SESSION_KEY + 1)
    if rstrtmgr.RmStartSession(ctypes.byref(sesion), 0, clave) != 0:
        return []
    try:
        archivos = (wintypes.LPCWSTR * 1)(str(ruta))
        if rstrtmgr.RmRegisterResources(sesion, 1, archivos, 0, None, 0, None) != 0:
            return []

        necesarios = ctypes.c_uint(0)
        cuantos = ctypes.c_uint(0)
        razones = wintypes.DWORD(0)
        rc = rstrtmgr.RmGetList(sesion, ctypes.byref(necesarios),
                                ctypes.byref(cuantos), None, ctypes.byref(razones))
        if rc not in (0, ERROR_MORE_DATA) or necesarios.value == 0:
            return []

        n = necesarios.value
        info = (RM_PROCESS_INFO * n)()
        cuantos = ctypes.c_uint(n)
        if rstrtmgr.RmGetList(sesion, ctypes.byref(necesarios),
                              ctypes.byref(cuantos), info, ctypes.byref(razones)) != 0:
            return []

        salida = []
        for i in range(cuantos.value):
            p = info[i]
            salida.append({
                "pid": p.Process.dwProcessId,
                "app": (p.strAppName or "").strip() or "(aplicacion desconocida)",
                "servicio": (p.strServiceShortName or "").strip(),
            })
        return salida
    finally:
        rstrtmgr.RmEndSession(sesion)


def quien_bloquea(ruta):
    """Lineas explicativas de quien tiene el archivo. Nunca lanza excepcion."""
    ruta = Path(ruta)
    lineas = []
    try:
        procesos = _procesos_que_bloquean(ruta)
    except Exception:
        procesos = []
    try:
        duenos = _duenos_office(ruta)
    except Exception:
        duenos = []

    for p in procesos:
        detalle = f"PID {p['pid']}"
        if p["servicio"]:
            detalle += f", servicio {p['servicio']}"
        lineas.append(f"{p['app']}   ({detalle})   en este equipo")
    for d in duenos:
        lineas.append(f"Figura como abierto por: {d}")
    if not lineas and (ruta.parent / f"~${ruta.name}").exists():
        lineas.append("Hay un archivo de bloqueo de Office (~$) junto al documento.")
    return lineas


def comprobar_escribible(ruta):
    """Aborta ANTES de tocar nada si el documento está en uso.

    Word mantiene el .docx abierto mientras lo tiene en pantalla. Si se
    escribe encima en ese momento, el archivo queda inservible (bytes en
    cero). Antes esto solo se avisaba en la documentación; ahora se impide.
    """
    ruta = Path(ruta)
    if not ruta.exists():
        return
    bloqueo = ruta.parent / f"~${ruta.name}"
    abierto_en_office = bloqueo.exists()
    try:
        with open(ruta, "r+b"):
            pass
        if not abierto_en_office:
            return
    except PermissionError:
        pass
    except OSError as e:
        raise ValueError(f"No puedo escribir en el documento:\n  {ruta}\n  {e}")

    culpables = quien_bloquea(ruta)
    if culpables:
        detalle = "\n\nQuién lo tiene:\n" + "\n".join(f"  - {c}" for c in culpables)
        if not any("en este equipo" in c for c in culpables):
            detalle += (
                "\n\nNinguna aplicación de ESTE equipo lo retiene: es probable\n"
                "que lo tenga abierto otra persona a través de OneDrive."
            )
    else:
        detalle = (
            "\n\nNo pude identificar qué lo retiene. Suele ser Word en este\n"
            "equipo, o una sincronización de OneDrive en curso."
        )

    raise ValueError(
        f"El documento está abierto:\n"
        f"  archivo:  {ruta.name}\n"
        f"  carpeta:  {ruta.parent}"
        f"{detalle}\n\n"
        "Ciérrelo y vuelva a ejecutar.\n\n"
        "No se ha modificado nada. Escribir sobre un documento que Word\n"
        "tiene abierto lo deja inservible, así que la operación se detiene\n"
        "aquí a propósito."
    )


def guardar_seguro(doc, ruta):
    """Guarda SIN cambiar la identidad del archivo en el disco.

    Antes esto era doc.save() a un temporal de la misma carpeta y luego
    os.replace(). Es atómico y en un disco normal está bien, pero en una
    carpeta de OneDrive rompe la sincronización, y de forma silenciosa:

      - os.replace() borra el archivo original y pone otro en su sitio. El
        archivo que queda tiene un File ID de NTFS NUEVO.
      - OneDrive lleva su base de datos indexada por ese File ID, no por la
        ruta. Al no reconocerlo, no lo lee como «el documento cambió», sino
        como «el documento desapareció y hay uno desconocido en su sitio».
      - Su forma de resolver ese conflicto es reponer la versión que tiene
        en el servidor. Minutos después el archivo local vuelve a ser el de
        antes y los cambios se han perdido sin un solo mensaje de error.

    Con Archivos a Petición es todavía más claro: el original es un punto
    de reanalisis (placeholder) y el temporal no, así que ni siquiera son
    el mismo tipo de archivo.

    Ahora se hace al revés, que es como escribe Word: se serializa entero a
    un temporal FUERA de la carpeta sincronizada (para que OneDrive no vea
    aparecer y desaparecer archivos sueltos), se comprueba que el resultado
    es un .docx legible, y solo entonces se vuelca sobre el archivo original
    abriéndolo en modo r+b. El archivo conserva su identidad y OneDrive lo
    ve como lo que es: una modificación normal.
    """
    import tempfile
    import zipfile

    ruta = Path(ruta)

    # 1. Serializar fuera de la carpeta sincronizada.
    tmp = Path(tempfile.gettempdir()) / f"fs_{os.getpid()}_{ruta.name}"
    try:
        doc.save(str(tmp))
        if not zipfile.is_zipfile(tmp):
            raise ValueError(
                "El documento generado no es un .docx legible; no se "
                "escribe nada sobre el original."
            )
        datos = tmp.read_bytes()
    finally:
        try:
            tmp.unlink()
        except OSError:
            pass

    # 2. Volcar sobre el original conservando su identidad.
    if ruta.exists():
        with open(ruta, "r+b") as f:
            f.write(datos)
            f.truncate()
            f.flush()
            os.fsync(f.fileno())
    else:
        ruta.write_bytes(datos)


# --------------------------------------------------------------------------- #
#  Tablas
# --------------------------------------------------------------------------- #
def _anchos_de(tbl):
    """Lee los anchos del w:tblGrid. Si el usuario arrastra las columnas en
    Word, el tblGrid se actualiza y el refresco respeta la nueva medida."""
    grid = tbl.find(qn("w:tblGrid")) if tbl is not None else None
    if grid is None:
        return list(ANCHOS_DEFECTO)
    anchos = []
    for gc in grid.findall(qn("w:gridCol")):
        try:
            anchos.append(int(gc.get(qn("w:w"))))
        except (TypeError, ValueError):
            anchos.append(ANCHOS_DEFECTO[len(anchos) % 4])
    return anchos or list(ANCHOS_DEFECTO)


def _tblPr_defecto():
    tblPr = OxmlElement("w:tblPr")
    w = OxmlElement("w:tblW")
    w.set(qn("w:w"), "0")
    w.set(qn("w:type"), "auto")
    tblPr.append(w)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    return tblPr


def _tblGrid(anchos):
    grid = OxmlElement("w:tblGrid")
    for a in anchos:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(a))
        grid.append(gc)
    return grid


def _fila_encabezado(ctx, anchos):
    tr = OxmlElement("w:tr")
    trPr = OxmlElement("w:trPr")
    trPr.append(OxmlElement("w:tblHeader"))
    tr.append(trPr)

    fa = str(ctx.get("fecha_actual", "") or "")
    fp = str(ctx.get("fecha_previa", "") or "")
    ea = str(ctx.get("estado_actual", "") or "")
    ep = str(ctx.get("estado_previo", "") or "")

    def cab(texto_, sub, alineacion):
        ps = [_parrafo_texto(texto_, negrita=True, alineacion=alineacion)]
        if sub:
            ps.append(_parrafo_texto(sub, negrita=False, cursiva=True, alineacion=alineacion))
        return ps

    tr.append(_celda(anchos[0], cab("", "", None), borde_inf="single"))
    tr.append(_celda(anchos[1], cab("Nota", "", "center"), borde_inf="single"))
    tr.append(_celda(anchos[2], cab(fa, ea, "right"), borde_inf="single"))
    tr.append(_celda(anchos[3], cab(fp, ep, "right"), borde_inf="single"))
    return tr


def _fila_de_linea(linea, anchos):
    tipo = linea.get("tipo", "I")
    est = ESTILO_ETIQUETA.get(tipo, ESTILO_ETIQUETA["I"])
    tr = OxmlElement("w:tr")

    borde_sup = "single" if tipo in ("S", "T") else None
    borde_inf = "double" if tipo == "T" else None

    etiqueta = str(linea.get("etiqueta", "") or "")
    nota = str(linea.get("nota", "") or "")
    actual = str(linea.get("actual", "") or "")
    previo = str(linea.get("previo", "") or "")

    if tipo in ("H", "N"):
        nota = actual = previo = ""

    p_et = _parrafo_texto(
        etiqueta,
        negrita=est["negrita"],
        cursiva=est.get("cursiva", False),
        sangria=est["sangria"],
        espacio_antes=est["espacio_antes"],
    )
    tr.append(_celda(anchos[0], [p_et], borde_sup=borde_sup, borde_inf=borde_inf))
    tr.append(
        _celda(
            anchos[1],
            [_parrafo_texto(nota, alineacion="center", espacio_antes=est["espacio_antes"])],
            borde_sup=borde_sup,
            borde_inf=borde_inf,
        )
    )
    for ancho, val in ((anchos[2], actual), (anchos[3], previo)):
        tr.append(
            _celda(
                ancho,
                [
                    _parrafo_texto(
                        val,
                        negrita=est["negrita"],
                        alineacion="right",
                        espacio_antes=est["espacio_antes"],
                    )
                ],
                borde_sup=borde_sup,
                borde_inf=borde_inf,
            )
        )
    return tr


def _tabla(ctx, lineas, anchos=None, tblPr=None):
    anchos = list(anchos or ANCHOS_DEFECTO)
    while len(anchos) < 4:
        anchos.append(ANCHOS_DEFECTO[len(anchos)])
    tbl = OxmlElement("w:tbl")
    tbl.append(copy.deepcopy(tblPr) if tblPr is not None else _tblPr_defecto())
    tbl.append(_tblGrid(anchos[:4]))
    tbl.append(_fila_encabezado(ctx, anchos))
    for linea in lineas:
        tbl.append(_fila_de_linea(linea, anchos))
    return tbl


def _lineas_de_tabla(nombre, ctx):
    """Qué líneas del Excel alimentan la tabla `fs-tabla-<nombre>`.

    'principal' -> todas.
    Cualquier otro nombre -> las líneas de la sección cuyo encabezado (fila
    de tipo H) produce esa misma clave. Permite partir el estado en varias
    tablas sin tocar el código: basta con nombrar el control de contenido
    'fs-tabla-current_assets' y esa tabla recibe solo esa sección.
    """
    lineas = ctx.get("lineas", [])
    if nombre == C.TABLA_PRINCIPAL:
        return list(lineas)

    seleccion = []
    dentro = False
    for linea in lineas:
        if linea.get("tipo") == "H":
            dentro = C.clave(linea.get("etiqueta")) == nombre
            if dentro:
                seleccion.append(linea)
            continue
        if dentro:
            seleccion.append(linea)
    return seleccion


# --------------------------------------------------------------------------- #
#  Construcción / reparación del andamiaje
# --------------------------------------------------------------------------- #
def _bloque_prosa(nombre, texto_guia):
    sdt = _sdt(
        C.tag_prosa(nombre),
        alias=f"Redacción — {nombre}",
        bloqueado=False,          # ESTA es la zona que la persona edita
    )
    cont = _contenido(sdt)
    cont.append(_parrafo_texto(texto_guia))
    # Un párrafo vacío de holgura. Con el documento protegido, el rango
    # editable termina en la última marca de párrafo: sin este hueco, quien
    # solo tiene el rol Redactor puede corregir el texto existente pero no
    # empezar un párrafo nuevo al final de la zona.
    cont.append(_parrafo())
    return sdt


def _bloque_campo(nombre, valor=""):
    sdt = _sdt(C.tag_campo(nombre), alias=f"Campo — {nombre}", en_linea=True)
    _contenido(sdt).append(_run(valor or f"«{nombre}»"))
    return sdt


def _salto_de_pagina():
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    p.append(r)
    return p


#: Encabezado del apartado que se añade a un documento que YA tenia
#: redaccion. Sin el, las cifras aparecian pegadas al final del texto de la
#: persona, sin que nada dijera donde empieza lo que mantiene el programa.
TITULO_APARTADO = "Estados financieros"


def construir(ruta, ctx=None, verbose=True, cfg_bitacora=None, apartado=None):
    """Añade al documento lo que le falte para cumplir el contrato.

    Idempotente: correrla dos veces no duplica nada. No borra prosa ni
    reordena lo que ya exista; solo agrega las anclas ausentes al final.

    `apartado`: si el documento ya trae redacción propia, lo que se añade va
    detrás de un salto de página y un título, para que se lea como una
    sección aparte y no como una continuación del texto de la persona. None
    = decidirlo mirando el documento; True/False = forzarlo.
    """
    comprobar_escribible(ruta)
    doc = Document(str(ruta))
    idx = _indexar(doc)
    añadidos = []
    ctx = ctx or {}

    # Qué había ANTES de tocar nada, y dónde acababa. Las dos cosas hacen
    # falta para poder meter después el encabezado del apartado justo
    # delante de lo primero que se añada.
    parrafos_previos, tablas_previas = _contenido_visible(doc)
    if apartado is None:
        apartado = parrafos_previos > _UMBRAL_EN_BLANCO or tablas_previas > 0
    # La lista se GUARDA, no se convierte en un conjunto de id(): mientras
    # haya una referencia viva, lxml devuelve siempre el mismo objeto para
    # el mismo nodo y la comparacion por identidad vale. Sin la referencia,
    # los envoltorios se reciclan y las direcciones se repiten.
    ya_estaban = list(doc.element.body)

    def falta(tag):
        return tag not in idx

    # --- encabezado del estado -------------------------------------------- #
    #  Cada campo va dentro de una frase con sentido, no en una línea suelta:
    #  el documento tiene que poder imprimirse tal cual desde el primer día.
    #  Un campo que ya exista en el documento no se vuelve a poner, así que
    #  la persona puede moverlos donde quiera y 'reparar' los respeta.
    def linea_campos(piezas, **kw):
        """piezas: lista de str (texto fijo) o ('campo', nombre)."""
        nombres = [x[1] for x in piezas if isinstance(x, tuple)]
        if not any(falta(C.tag_campo(n)) for n in nombres):
            return                      # todos sus campos ya existen: no repetir
        p = _parrafo(**kw)
        for pieza in piezas:
            if isinstance(pieza, tuple):
                nombre = pieza[1]
                if falta(C.tag_campo(nombre)):
                    p.append(_bloque_campo(nombre, str(ctx.get(nombre, ""))))
                    añadidos.append(C.tag_campo(nombre))
                else:
                    p.append(_run(str(ctx.get(nombre, "")) or f"«{nombre}»"))
            else:
                p.append(_run(pieza))
        _cuerpo_append(doc, p)

    linea_campos([("campo", "empresa")], alineacion="center", espacio_antes=240)
    linea_campos([("campo", "titulo")], alineacion="center")
    linea_campos(
        ["Al ", ("campo", "fecha_actual"), " (", ("campo", "estado_actual"),
         ") — comparado con ", ("campo", "fecha_previa"),
         " (", ("campo", "estado_previo"), ")"],
        alineacion="center",
    )
    linea_campos(
        ["Cifras expresadas en ", ("campo", "moneda"), ", en unidades de ",
         ("campo", "miles")],
        alineacion="center",
    )

    # --- zonas de prosa + tabla ------------------------------------------- #
    if falta(C.tag_prosa("introduccion")):
        _cuerpo_append(
            doc,
            _bloque_prosa(
                "introduccion",
                "Zona de redacción libre. Escriba aquí lo que quiera: este "
                "texto NO se toca al refrescar las cifras.",
            ),
        )
        añadidos.append(C.tag_prosa("introduccion"))

    tag_tabla = C.tag_tabla(C.TABLA_PRINCIPAL)
    if falta(tag_tabla):
        sdt = _sdt(tag_tabla, alias="Tabla — estado principal")
        lineas = _lineas_de_tabla(C.TABLA_PRINCIPAL, ctx) if ctx else []
        _contenido(sdt).append(_tabla(ctx, lineas))
        _cuerpo_append(doc, sdt)
        # Word necesita un párrafo tras una tabla al final del cuerpo
        _cuerpo_append(doc, _parrafo())
        añadidos.append(tag_tabla)

    if falta(C.tag_prosa("analisis")):
        _cuerpo_append(
            doc,
            _bloque_prosa(
                "analisis",
                "Zona de análisis. Aquí puede intercalar cifras vivas: use "
                "«fs_documento.py catalogo» para ver las claves disponibles.",
            ),
        )
        añadidos.append(C.tag_prosa("analisis"))

    # --- bitácora y foto -------------------------------------------------- #
    # La bitácora solo se crea DENTRO del documento si se pide expresamente.
    # Por defecto vive en un .log aparte para no estorbar la redacción.
    modo_bitacora = str((cfg_bitacora or {}).get("bitacora", "archivo")).lower()
    quiere_registro = modo_bitacora in ("documento", "ambos")
    if quiere_registro and falta(C.TAG_REGISTRO):
        sdt = _sdt(C.TAG_REGISTRO, alias="Bitácora de actualizaciones", bloqueado=True)
        _contenido(sdt).append(_parrafo_texto("Bitácora de actualizaciones", negrita=True))
        _cuerpo_append(doc, sdt)
        añadidos.append(C.TAG_REGISTRO)

    if falta(C.TAG_META):
        sdt = _sdt(C.TAG_META, alias="Metadatos (oculto)", bloqueado=True)
        _contenido(sdt).append(_parrafo_texto("{}", oculto=True))
        _cuerpo_append(doc, sdt)
        añadidos.append(C.TAG_META)

    # La separación del apartado va DELANTE de lo primero que se haya
    # añadido: no al final del cuerpo, que ahí ya está la tabla.
    #
    # Solo el salto de página: el propio andamiaje empieza por las líneas
    # centradas de empresa y título, que ya hacen de encabezado. Poner otro
    # encima lo diría dos veces. Si esas líneas ya existían en el documento
    # (alguien las movió), entonces sí hace falta un rótulo.
    if apartado and añadidos:
        nuevos = [el for el in doc.element.body
                  if not any(el is viejo for viejo in ya_estaban)]
        if nuevos:
            nuevos[0].addprevious(_salto_de_pagina())
            if C.tag_campo("titulo") not in añadidos:
                nuevos[0].addprevious(
                    _parrafo_texto(TITULO_APARTADO, negrita=True,
                                   alineacion="center", espacio_antes=240))

    # Aspecto: sin recuadro en las zonas de redacción ni en los metadatos,
    # para que escribir se parezca a escribir en un documento normal.
    normalizar_apariencia(
        doc, datos=str((cfg_bitacora or {}).get("apariencia_datos", "boundingBox")))
    _sdt_meta = _indexar(doc).get(C.TAG_META)
    if _sdt_meta is not None:
        for _p in _contenido(_sdt_meta).findall(qn("w:p")):
            _ocultar_parrafo(_p)

    guardar_seguro(doc, ruta)
    if verbose:
        if añadidos:
            print(f"  Anclas añadidas ({len(añadidos)}):")
            for t in añadidos:
                print(f"    + {t}")
        else:
            print("  Nada que añadir: el documento ya cumple el contrato.")
    return añadidos


def preparar(ruta, ctx=None, cfg=None, verbose=True, respaldar=True):
    """Deja un documento cualquiera listo para el refresco, y dice cómo.

    Es `construir` con la copia de seguridad delante y un informe legible
    detrás. Devuelve (estado_previo, añadidos): el estado dice si el
    documento ya venía integrado, estaba en blanco o traía redacción propia,
    que es justo lo que el usuario necesita leer antes de aceptar.

    `respaldar=False` para cuando el llamante ya hizo la copia (el refresco
    la hace él): un segundo .bak machacaría el bueno con el ya modificado.
    """
    estado_, _familias, detalle = clasificar_documento(ruta)
    if estado_ == LISTO:
        if verbose:
            print("  Ya está integrado: no hay que añadirle nada.")
        return estado_, []

    if respaldar:
        _respaldar(ruta)
    if verbose:
        if estado_ == EN_BLANCO:
            print("  El documento está en blanco: se usa de base.")
        else:
            print(f"  El documento trae {detalle['parrafos']} párrafos y "
                  f"{detalle['tablas']} tablas escritas.")
            print("  Se le añade el estado como apartado aparte; su texto no "
                  "se toca.")
    añadidos = construir(ruta, ctx, verbose=verbose, cfg_bitacora=cfg,
                         apartado=(estado_ == CON_TEXTO))
    return estado_, añadidos


def crear_base(destino, ctx=None, cfg=None, verbose=True):
    """Crea desde cero un .docx con TODAS las regiones, y las cifras dentro.

    Es la contraparte de generador_fs: aquel renderiza una plantilla de
    Word y produce un documento sin regiones —una foto, imposible de
    refrescar después—; este produce el documento vivo, el que la opción de
    actualizar sabe mantener al día.

    `destino` puede estar en el disco local o dentro de OneDrive: es la
    misma operación.
    """
    destino = Path(destino)
    if destino.exists():
        comprobar_escribible(destino)
    destino.parent.mkdir(parents=True, exist_ok=True)

    # Un Word en blanco de verdad: la plantilla por defecto de python-docx
    # trae los estilos normales de Office y ni un párrafo con texto.
    doc = Document()
    guardar_seguro(doc, destino)

    # apartado=False: no hay nada delante de lo que separarse, así que
    # tampoco hace falta el salto de página.
    construir(destino, ctx, verbose=verbose, cfg_bitacora=cfg, apartado=False)
    return destino


# --------------------------------------------------------------------------- #
#  Refresco en el sitio
# --------------------------------------------------------------------------- #
def _escribir_en_linea(sdt, valor):
    """Reemplaza el texto de un control en línea conservando su formato."""
    cont = _contenido(sdt)
    rPr = None
    for r in cont.iter(qn("w:r")):
        hallado = r.find(qn("w:rPr"))
        if hallado is not None:
            rPr = copy.deepcopy(hallado)
        break
    _vaciar(cont)
    r = OxmlElement("w:r")
    if rPr is not None:
        r.append(rPr)
    t = OxmlElement("w:t")
    t.text = "" if valor is None else str(valor)
    t.set(XMLSPACE, "preserve")
    r.append(t)
    cont.append(r)


def _leer_en_linea(sdt):
    return "".join(t.text or "" for t in _contenido(sdt).iter(qn("w:t")))


def _leer_meta(idx):
    sdt = idx.get(C.TAG_META)
    if sdt is None:
        return {}
    crudo = "".join(t.text or "" for t in _contenido(sdt).iter(qn("w:t")))
    try:
        return json.loads(crudo) or {}
    except (json.JSONDecodeError, TypeError):
        return {}


def _guardar_meta(idx, ctx, origen):
    sdt = idx.get(C.TAG_META)
    if sdt is None:
        return
    foto = {
        "fecha": datetime.now().isoformat(timespec="seconds"),
        "origen": origen,
        "lineas": [
            {
                "e": l.get("etiqueta", ""),
                "t": l.get("tipo", ""),
                "a": l.get("actual", ""),
                "p": l.get("previo", ""),
            }
            for l in ctx.get("lineas", [])
        ],
    }
    cont = _contenido(sdt)
    _vaciar(cont)
    cont.append(
        _parrafo_texto(json.dumps(foto, ensure_ascii=False, separators=(",", ":")), oculto=True)
    )


def _claves_diff(filas, etiqueta_de, tipo_de):
    """Un nombre estable y legible para cada fila, apto para comparar.

    Las filas de subtotal (S) no traen etiqueta: todas se llamarían igual y
    el diff las daría por nuevas en cada corrida. Se las nombra por la
    sección en la que caen ('Subtotal de Current assets:'), y si aun así
    dos coinciden se numeran.
    """
    nombres = []
    seccion = ""
    vistos = {}
    for f in filas:
        etiqueta = (etiqueta_de(f) or "").strip()
        tipo = tipo_de(f) or ""
        if tipo == "H" and etiqueta:
            seccion = etiqueta.rstrip(":").strip()
        if etiqueta:
            base = etiqueta
        elif tipo == "S":
            base = f"Subtotal de {seccion}" if seccion else "Subtotal"
        else:
            base = f"(sin etiqueta {tipo})"
        vistos[base] = vistos.get(base, 0) + 1
        nombres.append(base if vistos[base] == 1 else f"{base} #{vistos[base]}")
    return nombres


def _calcular_cambios(meta_previa, ctx):
    previas = meta_previa.get("lineas") or []
    if not previas:
        return ["Primera actualización: no hay versión anterior con la que comparar."]

    nom_antes = _claves_diff(previas, lambda d: d.get("e"), lambda d: d.get("t"))
    nom_ahora = _claves_diff(
        ctx.get("lineas", []), lambda d: d.get("etiqueta"), lambda d: d.get("tipo")
    )
    antes = {n: f for n, f in zip(nom_antes, previas)}

    cambios = []
    for nombre, l in zip(nom_ahora, ctx.get("lineas", [])):
        a = antes.pop(nombre, None)
        if a is None:
            cambios.append(f"Nueva fila: {nombre}  {l.get('actual') or '—'}")
        elif a.get("a") != l.get("actual") or a.get("p") != l.get("previo"):
            cambios.append(
                f"{nombre}: {a.get('a') or '—'} → {l.get('actual') or '—'}"
                f"  (comparativo {a.get('p') or '—'} → {l.get('previo') or '—'})"
            )
    for nombre in antes:
        cambios.append(f"Fila retirada: {nombre}")
    return cambios or ["Sin cambios en las cifras respecto de la última actualización."]


def ruta_bitacora(cfg, documento):
    """Dónde se escribe el registro de actualizaciones.

    Por defecto FUERA del documento: un .log junto al ejecutable, en
    salidas\\. El documento de OneDrive tiene que quedar limpio para
    redactar; una bitácora creciendo dentro estorba la lectura y provoca
    conflictos de sincronización cuando dos personas lo abren.
    """
    import generador_fs as G

    crudo = str(cfg.get("bitacora_archivo") or "").strip()
    if crudo:
        ruta = Path(crudo)
        return ruta if ruta.is_absolute() else (G.BASE / ruta)
    nombre = C.clave(Path(documento).stem)[:40] or "documento"
    return G.BASE / "salidas" / f"bitacora_{nombre}.log"


def escribir_bitacora_archivo(ruta_log, documento, cambios, origen, informe=None):
    """Añade una entrada al final del .log. Nunca lo reescribe."""
    ruta_log = Path(ruta_log)
    ruta_log.parent.mkdir(parents=True, exist_ok=True)
    sello = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    lineas = [
        "=" * 72,
        f"{sello}   {Path(documento).name}",
        f"  origen: {origen}",
    ]
    if informe:
        filas = sum(n for _, n in informe.get("tablas", []))
        lineas.append(
            f"  escrito: {filas} filas de tabla, {informe.get('campos', 0)} campos, "
            f"{informe.get('datos', 0)} cifras en el texto"
        )
        if informe.get("huerfanos"):
            lineas.append(f"  anclas sin dato en el Excel: "
                          f"{', '.join(informe['huerfanos'])}")
    lineas.append("  cambios:")
    for c in cambios:
        lineas.append(f"    - {c}")
    lineas.append("")

    with ruta_log.open("a", encoding="utf-8") as f:
        f.write("\n".join(lineas) + "\n")
    return ruta_log


def quitar_registro_del_documento(ruta, verbose=True):
    """Elimina del .docx la región de bitácora y todo lo que llevara dentro.

    Se usa al pasar la bitácora a un archivo aparte: deja de tener sentido
    arrastrar el histórico dentro del documento de trabajo.
    """
    comprobar_escribible(ruta)
    doc = Document(str(ruta))
    quitados = 0
    for sdt in list(doc.element.body.iter(qn("w:sdt"))):
        if _tag_de(sdt) == C.TAG_REGISTRO:
            padre = sdt.getparent()
            if padre is not None:
                padre.remove(sdt)
                quitados += 1
    if quitados:
        guardar_seguro(doc, ruta)
    if verbose:
        print(f"  Bloques de bitácora retirados del documento: {quitados}")
    return quitados
def _escribir_registro(idx, cambios, origen):
    sdt = idx.get(C.TAG_REGISTRO)
    if sdt is None:
        return False
    cont = _contenido(sdt)
    sello = datetime.now().strftime("%Y-%m-%d %H:%M")
    bloque = [
        _parrafo_texto(f"Actualización {sello} — origen: {origen}", negrita=True,
                       espacio_antes=120)
    ]
    for c in cambios[:40]:
        bloque.append(_parrafo_texto(f"• {c}", sangria=220))
    if len(cambios) > 40:
        bloque.append(_parrafo_texto(f"• … y {len(cambios) - 40} cambios más.", sangria=220))

    titulo = cont.find(qn("w:p"))
    ancla = titulo if titulo is not None else None
    for el in reversed(bloque):
        if ancla is not None:
            ancla.addnext(el)
        else:
            cont.append(el)
    return True


def insertar_dato(ruta, clave_, campo, zona="analisis", antes="", despues="",
                  valor_inicial="—", verbose=True):
    """Añade un párrafo con una cifra viva al final de una zona de prosa.

    Es el equivalente por línea de órdenes al botón «insertar dato» del
    add-in: deja el control de contenido en línea, bloqueado y con la
    etiqueta correcta, listo para que el siguiente refresco lo rellene.
    """
    if campo not in C.CAMPOS_DATO:
        raise ValueError(
            f"'{campo}' no es un campo válido.\n"
            f"Válidos: {', '.join(C.CAMPOS_DATO)}"
        )
    comprobar_escribible(ruta)
    doc = Document(str(ruta))
    idx = _indexar(doc)
    sdt_zona = idx.get(C.tag_prosa(zona))
    if sdt_zona is None:
        raise ValueError(
            f"No existe la zona de prosa 'fs-prosa-{zona}' en el documento.\n"
            f"Zonas disponibles: "
            f"{', '.join(sorted(n for t in idx for f, n, _ in [C.descomponer(t)] if f == C.FAM_PROSA)) or '(ninguna)'}"
        )

    tag = C.tag_dato(clave_, campo)
    if tag in idx:
        if verbose:
            print(f"  El ancla {tag} ya existe en el documento; no se duplica.")
        return False

    p = _parrafo()
    if antes:
        p.append(_run(antes))
    sdt = _sdt(tag, alias=f"{clave_} ({campo})", en_linea=True)
    _contenido(sdt).append(_run(valor_inicial))
    p.append(sdt)
    if despues:
        p.append(_run(despues))
    _contenido(sdt_zona).append(p)
    guardar_seguro(doc, ruta)
    if verbose:
        print(f"  + {tag}  en la zona '{zona}'")
    return True


def refrescar(ruta, ctx, origen="", con_registro=True, verbose=True,
              cfg=None):
    """Reescribe SOLO las regiones de datos. Devuelve un informe."""
    comprobar_escribible(ruta)
    doc = Document(str(ruta))
    idx = _indexar(doc)
    valores, colisiones = C.construir_valores(ctx)

    meta_previa = _leer_meta(idx)
    cambios = _calcular_cambios(meta_previa, ctx)

    informe = {
        "tablas": [], "campos": 0, "datos": 0,
        "huerfanos": [], "colisiones": colisiones,
        "cambios": cambios, "sin_ancla_prosa": 0,
    }

    for tag, sdt in idx.items():
        familia, nombre, campo = C.descomponer(tag)

        if familia == C.FAM_TABLA:
            cont = _contenido(sdt)
            vieja = cont.find(qn("w:tbl"))
            anchos = _anchos_de(vieja)
            tblPr = vieja.find(qn("w:tblPr")) if vieja is not None else None
            lineas = _lineas_de_tabla(nombre, ctx)
            nueva = _tabla(ctx, lineas, anchos=anchos, tblPr=tblPr)
            _vaciar(cont)
            cont.append(nueva)
            informe["tablas"].append((nombre, len(lineas)))

        elif familia == C.FAM_CAMPO:
            if tag in valores:
                _escribir_en_linea(sdt, valores[tag])
                informe["campos"] += 1
            else:
                informe["huerfanos"].append(tag)

        elif familia == C.FAM_DATO:
            if tag in valores:
                _escribir_en_linea(sdt, valores[tag])
                informe["datos"] += 1
            else:
                informe["huerfanos"].append(tag)

        elif familia == C.FAM_PROSA:
            informe["sin_ancla_prosa"] += 1

    # La bitácora va FUERA del documento salvo que se pida lo contrario:
    # un histórico creciendo dentro estorba a quien redacta y multiplica
    # los conflictos de sincronización en OneDrive.
    destino = str((cfg or {}).get("bitacora", "archivo")).lower()
    informe["con_registro"] = False
    informe["bitacora_archivo"] = None
    if con_registro and destino in ("documento", "ambos"):
        informe["con_registro"] = _escribir_registro(idx, cambios, origen)
    if con_registro and destino in ("archivo", "ambos"):
        informe["bitacora_archivo"] = escribir_bitacora_archivo(
            ruta_bitacora(cfg or {}, ruta), ruta, cambios, origen, informe)

    _guardar_meta(idx, ctx, origen)
    guardar_seguro(doc, ruta)
    return informe


# --------------------------------------------------------------------------- #
#  Verificación
# --------------------------------------------------------------------------- #
def verificar(ruta, ctx=None):
    doc = Document(str(ruta))
    idx = _indexar(doc)
    valores = C.construir_valores(ctx)[0] if ctx else {}

    rep = {
        "tablas": [], "campos": [], "datos": [], "prosa": [],
        "registro": C.TAG_REGISTRO in idx, "meta": C.TAG_META in idx,
        "desconocidos": [], "huerfanos": [], "sin_usar": [],
    }

    for tag, sdt in idx.items():
        familia, nombre, campo = C.descomponer(tag)
        if familia == C.FAM_TABLA:
            cont = _contenido(sdt)
            tiene = cont.find(qn("w:tbl")) is not None
            rep["tablas"].append((nombre, tiene))
        elif familia == C.FAM_CAMPO:
            rep["campos"].append((nombre, _leer_en_linea(sdt)))
            if valores and tag not in valores:
                rep["huerfanos"].append(tag)
        elif familia == C.FAM_DATO:
            rep["datos"].append((nombre, campo, _leer_en_linea(sdt)))
            if valores and tag not in valores:
                rep["huerfanos"].append(tag)
        elif familia == C.FAM_PROSA:
            rep["prosa"].append(nombre)
        elif familia in ("registro", "meta"):
            pass
        else:
            rep["desconocidos"].append(tag)

    if valores:
        usados = set(idx)
        rep["sin_usar"] = sorted(
            t for t in valores
            if t.startswith("fs-dato-") and t not in usados and valores[t]
        )
    return rep


# --------------------------------------------------------------------------- #
#  Apariencia de las regiones
# --------------------------------------------------------------------------- #
#  Word dibuja un recuadro gris alrededor de cada control de contenido. Para
#  la tabla y las cifras eso es útil: se ve de un vistazo qué lo mantiene el
#  Excel. Para las zonas de redacción y para los metadatos es un estorbo —
#  la persona quiere escribir sobre papel en blanco, no dentro de una caja.
#
#  w15:appearance="hidden" quita el recuadro sin quitar el ancla. Es de Word
#  2013 en adelante; las versiones viejas simplemente lo ignoran y siguen
#  mostrando el recuadro, que es un fallo inofensivo.
NS_W15 = "http://schemas.microsoft.com/office/word/2012/wordml"
nsmap.setdefault("w15", NS_W15)

#: Qué aspecto tiene cada familia. 'hidden' = sin recuadro.
#: Las zonas de redacción y los metadatos van SIEMPRE sin recuadro.
#: Las regiones de datos son configurables: con recuadro se ve de un vistazo
#: qué mantiene el Excel; sin él, la cifra se lee como una palabra más del
#: párrafo — sigue bloqueada y sigue refrescándose igual.
APARIENCIA_FIJA = {
    C.FAM_PROSA: "hidden",
    "registro": "hidden",
    "meta": "hidden",
}
FAMILIAS_CONFIGURABLES = (C.FAM_TABLA, C.FAM_CAMPO, C.FAM_DATO)


def apariencia_de(familia, datos="boundingBox"):
    if familia in APARIENCIA_FIJA:
        return APARIENCIA_FIJA[familia]
    if familia in FAMILIAS_CONFIGURABLES:
        return datos
    return None


def _poner_apariencia(pr, valor):
    """Fija w15:appearance dentro de un w:sdtPr, sin duplicarlo."""
    if pr is None:
        return
    for viejo in pr.findall(qn("w15:appearance")):
        pr.remove(viejo)
    ap = OxmlElement("w15:appearance")
    ap.set(qn("w15:val"), valor)
    # va al final del sdtPr, junto al resto de propiedades visuales
    pr.append(ap)


def normalizar_apariencia(doc, datos="boundingBox", verbose=False):
    """Ajusta el recuadro de todas las regiones. Idempotente.

    `datos` es el aspecto de las tablas, campos y cifras:
    'boundingBox' (recuadro visible) o 'hidden' (se lee como texto normal).
    """
    n = 0
    for sdt in doc.element.body.iter(qn("w:sdt")):
        familia, _, _ = C.descomponer(_tag_de(sdt))
        valor = apariencia_de(familia, datos)
        if valor is None:
            continue
        pr = sdt.find(qn("w:sdtPr"))
        actual = pr.find(qn("w15:appearance")) if pr is not None else None
        if actual is not None and actual.get(qn("w15:val")) == valor:
            continue
        _poner_apariencia(pr, valor)
        n += 1
    return n


def _ocultar_parrafo(p):
    """Marca el párrafo entero como texto oculto, incluida su marca final.

    Sin esto, un párrafo de metadatos con el texto oculto sigue ocupando una
    línea en blanco: lo que se ve en pantalla es un renglón vacío que nadie
    sabe de dónde sale.
    """
    pPr = p.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p.insert(0, pPr)
    rPr = pPr.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        pStyle = pPr.find(qn("w:pStyle"))
        if pStyle is not None:
            pStyle.addnext(rPr)
        else:
            pPr.insert(0, rPr)
    if rPr.find(qn("w:vanish")) is None:
        rPr.append(OxmlElement("w:vanish"))
    return p


def simplificar_documento(ruta, quitar_prosa=False, datos="boundingBox",
                          verbose=True):
    """Deja el documento cómodo para escribir.

    - quita el recuadro de las zonas de redacción y de los metadatos
    - oculta de verdad el párrafo de metadatos (no deja renglón en blanco)
    - opcionalmente disuelve las zonas de redacción, que dejan de ser
      controles y pasan a ser párrafos normales

    Las cifras y la tabla NO se tocan: su recuadro es informativo y su
    candado es lo que impide pisarlas.
    """
    comprobar_escribible(ruta)
    doc = Document(str(ruta))

    ajustadas = normalizar_apariencia(doc, datos=datos)

    ocultos = 0
    idx = _indexar(doc)
    sdt_meta = idx.get(C.TAG_META)
    if sdt_meta is not None:
        for p in _contenido(sdt_meta).findall(qn("w:p")):
            _ocultar_parrafo(p)
            ocultos += 1

    disueltas = []
    if quitar_prosa:
        for sdt in list(doc.element.body.iter(qn("w:sdt"))):
            tag = _tag_de(sdt)
            familia, nombre, _ = C.descomponer(tag)
            if familia != C.FAM_PROSA:
                continue
            cont = _contenido(sdt)
            padre = sdt.getparent()
            if cont is None or padre is None:
                continue
            for hijo in list(cont):
                sdt.addprevious(hijo)
            padre.remove(sdt)
            disueltas.append(tag)

    guardar_seguro(doc, ruta)
    if verbose:
        print(f"  Recuadros ajustados:        {ajustadas}")
        print(f"  Párrafos de metadatos ocultos: {ocultos}")
        if quitar_prosa:
            print(f"  Zonas de redacción disueltas: {len(disueltas)}")
            for t in disueltas:
                print(f"    - {t}   (ahora son párrafos normales)")
            print()
            print("  OJO: sin zonas de redacción no se puede usar el modo estricto")
            print("       de dos editores (proteger). Se pueden recrear con 'reparar'.")
    return ajustadas, ocultos, disueltas
# --------------------------------------------------------------------------- #
#  Candado por región: dejar (o no) que se escriba a mano en Word
# --------------------------------------------------------------------------- #
FAMILIAS_DATOS = (C.FAM_TABLA, C.FAM_CAMPO, C.FAM_DATO)


def cambiar_candado(ruta, bloquear=True, solo=None, verbose=True):
    """Pone o quita el candado de las regiones de datos.

    Con el candado puesto (por defecto) Word no deja teclear dentro de la
    tabla ni de las cifras. Sin candado, sí — pero OJO: lo que se escriba a
    mano lo machaca el siguiente refresco, porque la región sigue vinculada
    al Excel. Para conservar un valor escrito a mano hay que DESVINCULAR la
    región (ver desvincular_region).

    `solo` limita la operación a un tag exacto o a una clave.
    """
    comprobar_escribible(ruta)
    doc = Document(str(ruta))
    tocadas, nombres = 0, []

    for sdt in doc.element.body.iter(qn("w:sdt")):
        tag = _tag_de(sdt)
        familia, nombre, campo = C.descomponer(tag)
        if familia not in FAMILIAS_DATOS:
            continue
        if solo and solo not in (tag, nombre):
            continue

        pr = sdt.find(qn("w:sdtPr"))
        if pr is None:
            continue
        lock = pr.find(qn("w:lock"))

        if bloquear:
            if lock is None:
                lock = OxmlElement("w:lock")
                # w:lock va después de w:id y antes del tipo de control.
                tipo = pr.find(qn("w:text"))
                if tipo is None:
                    tipo = pr.find(qn("w:richText"))
                if tipo is not None:
                    tipo.addprevious(lock)
                else:
                    pr.append(lock)
            lock.set(qn("w:val"), "sdtContentLocked")
        elif lock is not None:
            pr.remove(lock)

        tocadas += 1
        nombres.append(tag)

    guardar_seguro(doc, ruta)
    if verbose:
        verbo = "bloqueadas" if bloquear else "desbloqueadas"
        print(f"  Regiones {verbo}: {tocadas}")
        if solo:
            for t in nombres:
                print(f"    - {t}")
    return tocadas


def estado_candado(ruta):
    """Como estan AHORA las cifras, sin modificar nada.

    Devuelve (bloqueadas, total, proteccion). `proteccion` es el modo de
    w:documentProtection ('readOnly', 'forms'...) o None si el documento no
    esta protegido.

    Existe porque las opciones «permitir editar» y «volver a proteger» eran
    dos botones ciegos: nadie decia en que estado estaba el documento, asi
    que pulsarlos no parecia tener efecto.
    """
    doc = Document(str(ruta))
    bloqueadas, total = _contar_candados(doc)
    prot = doc.settings.element.find(qn("w:documentProtection"))
    return bloqueadas, total, (prot.get(qn("w:edit")) if prot is not None else None)


def _contar_candados(doc):
    """(regiones de datos con candado, regiones de datos en total)."""
    bloqueadas = total = 0
    for sdt in doc.element.body.iter(qn("w:sdt")):
        familia, _, _ = C.descomponer(_tag_de(sdt))
        if familia not in FAMILIAS_DATOS:
            continue
        total += 1
        pr = sdt.find(qn("w:sdtPr"))
        lock = pr.find(qn("w:lock")) if pr is not None else None
        if lock is not None and lock.get(qn("w:val")) in (
                "sdtContentLocked", "sdtLocked"):
            bloqueadas += 1
    return bloqueadas, total


def desvincular_region(ruta, seleccion, verbose=True):
    """Convierte una región en texto normal: deja de refrescarse.

    Es la salida definitiva para una cifra que hay que escribir a mano y
    que debe SOBREVIVIR a los refrescos. El texto que hubiera dentro se
    conserva tal cual, pero ya no hay ancla: el motor no volverá a tocarlo
    ni lo reportará como huérfano.

    No tiene vuelta atrás automática; para volver a vincularla se usa
    'insertar'.
    """
    comprobar_escribible(ruta)
    doc = Document(str(ruta))
    quitadas = []

    for sdt in list(doc.element.body.iter(qn("w:sdt"))):
        tag = _tag_de(sdt)
        familia, nombre, campo = C.descomponer(tag)
        if familia not in FAMILIAS_DATOS:
            continue
        if seleccion not in (tag, nombre):
            continue

        cont = _contenido(sdt)
        padre = sdt.getparent()
        if cont is None or padre is None:
            continue
        # el contenido pasa a ocupar el sitio del control
        for hijo in list(cont):
            sdt.addprevious(hijo)
        padre.remove(sdt)
        quitadas.append(tag)

    if quitadas:
        guardar_seguro(doc, ruta)
    if verbose:
        if quitadas:
            print(f"  Regiones desvinculadas: {len(quitadas)}")
            for t in quitadas:
                print(f"    - {t}   (ahora es texto normal)")
        else:
            print(f"  No encontré ninguna región que coincida con '{seleccion}'.")
    return quitadas
# --------------------------------------------------------------------------- #
#  Protección (los dos editores)
# --------------------------------------------------------------------------- #
#  Orden de w:settings según ECMA-376; documentProtection debe ir en su sitio
#  o Word se queja al abrir.
_ORDEN_SETTINGS = [
    "writeProtection", "view", "zoom", "removePersonalInformation",
    "removeDateAndTime", "doNotDisplayPageBoundaries", "displayBackgroundShape",
    "printPostScriptOverText", "printFractionalCharacterWidth", "printFormsData",
    "embedTrueTypeFonts", "embedSystemFonts", "saveSubsetFonts", "saveFormsData",
    "mirrorMargins", "alignBordersAndEdges", "bordersDoNotSurroundHeader",
    "bordersDoNotSurroundFooter", "gutterAtTop", "hideSpellingErrors",
    "hideGrammaticalErrors", "activeWritingStyle", "proofState", "formsDesign",
    "attachedTemplate", "linkStyles", "stylePaneFormatFilter",
    "stylePaneSortMethod", "documentType", "mailMerge", "revisionView",
    "trackChanges", "doNotTrackMoves", "doNotTrackFormatting",
    "documentProtection",
]


def _hash_proteccion(clave_, salt, vueltas=100000):
    """Algoritmo de ECMA-376 para w:hash (SHA-1, cryptAlgorithmSid=4)."""
    h = hashlib.sha1(salt + clave_.encode("utf-16-le")).digest()
    for i in range(vueltas):
        h = hashlib.sha1(h + struct.pack("<I", i)).digest()
    return h


def _poner_documentProtection(doc, clave_):
    """Marca el documento como solo lectura, con la clave dada.

    Es la protección que Word IMPONE de verdad. El candado por región
    (w:lock) es una comodidad de la interfaz: Buscar y reemplazar lo
    atraviesa, y Word en el navegador ni lo mira.
    """
    settings = doc.settings.element
    for viejo in settings.findall(qn("w:documentProtection")):
        settings.remove(viejo)

    salt = os.urandom(16)
    prot = OxmlElement("w:documentProtection")
    prot.set(qn("w:edit"), "readOnly")
    prot.set(qn("w:enforcement"), "1")
    prot.set(qn("w:cryptProviderType"), "rsaFull")
    prot.set(qn("w:cryptAlgorithmClass"), "hash")
    prot.set(qn("w:cryptAlgorithmType"), "typeAny")
    prot.set(qn("w:cryptAlgorithmSid"), "4")
    prot.set(qn("w:cryptSpinCount"), "100000")
    prot.set(qn("w:hash"), base64.b64encode(_hash_proteccion(clave_, salt)).decode())
    prot.set(qn("w:salt"), base64.b64encode(salt).decode())

    pos = _ORDEN_SETTINGS.index("documentProtection")
    posteriores = set(_ORDEN_SETTINGS[pos + 1:])
    ancla = None
    for hijo in settings:
        nombre = hijo.tag.split("}")[-1]
        if nombre in posteriores or nombre not in _ORDEN_SETTINGS:
            ancla = hijo
            break
    if ancla is not None:
        ancla.addprevious(prot)
    else:
        settings.append(prot)



def proteger(ruta, clave_, verbose=True):
    """Deja el documento en solo lectura salvo las zonas fs-prosa-*.

    Rol REDACTOR  : abre y solo puede escribir dentro de las zonas de prosa.
    Rol EDITOR DE DATOS : conoce la clave (o usa el add-in), que desprotege,
                          refresca y vuelve a proteger.
    """
    comprobar_escribible(ruta)
    doc = Document(str(ruta))

    # 1. rangos editables alrededor de cada zona de prosa
    n = 0
    for sdt in doc.element.body.iter(qn("w:sdt")):
        tag = _tag_de(sdt)
        familia, _, _ = C.descomponer(tag)
        if familia != C.FAM_PROSA:
            continue
        cont = _contenido(sdt)
        if cont.find(qn("w:permStart")) is not None:
            continue
        n += 1
        ini = OxmlElement("w:permStart")
        ini.set(qn("w:id"), str(n))
        ini.set(qn("w:edGrp"), "everyone")
        fin = OxmlElement("w:permEnd")
        fin.set(qn("w:id"), str(n))
        cont.insert(0, ini)
        cont.append(fin)

    # 2. protección global
    _poner_documentProtection(doc, clave_)
    guardar_seguro(doc, ruta)
    if verbose:
        print(f"  Protegido. Zonas de prosa editables: {n}")
    return n


def _dentro_de_region(el):
    """¿Este bloque cuelga de algún control de contenido?"""
    nodo = el.getparent()
    while nodo is not None:
        if nodo.tag == qn("w:sdt"):
            return True
        nodo = nodo.getparent()
    return False


def _es_region_datos(el):
    if el.tag != qn("w:sdt"):
        return False
    familia, _, _ = C.descomponer(_tag_de(el))
    return familia in FAMILIAS_DATOS


def proteger_salvo_datos(ruta, clave_, verbose=True):
    """Solo lectura EXCEPTO todo lo que no sea una cifra.

    Es la vuelta del revés de proteger(): en vez de abrir huecos en las
    zonas de redacción, abre huecos en TODO menos en la tabla, los campos y
    las cifras intercaladas. Resultado: se escribe donde se quiera, y los
    números son intocables de verdad.

    Hace falta porque el candado de región (w:lock) es una comodidad de la
    interfaz, no una protección: Buscar y reemplazar lo atraviesa, y Word
    en el navegador ni lo mira. documentProtection sí lo impone Word.
    """
    comprobar_escribible(ruta)
    doc = Document(str(ruta))
    cuerpo = doc.element.body

    # quitar rangos editables anteriores para no acumularlos
    for viejo in list(cuerpo.iter(qn("w:permStart"))) + list(cuerpo.iter(qn("w:permEnd"))):
        padre = viejo.getparent()
        if padre is not None:
            padre.remove(viejo)

    sect = cuerpo.find(qn("w:sectPr"))

    # Un párrafo libre de cierre. Sin él, si el documento termina en la
    # tabla o en los metadatos, el último punto del documento cae fuera de
    # todo rango editable y el redactor no puede añadir un párrafo al final.
    hijos = [h for h in cuerpo if h is not sect]
    ultimo = hijos[-1] if hijos else None
    if ultimo is None or ultimo.tag != qn("w:p") or _dentro_de_region(ultimo):
        cierre = _parrafo()
        if sect is not None:
            sect.addprevious(cierre)
        else:
            cuerpo.append(cierre)

    hijos = [h for h in cuerpo if h is not sect]

    # agrupar tramos consecutivos que NO son regiones de datos
    tramos, actual = [], []
    for el in hijos:
        if _es_region_datos(el):
            if actual:
                tramos.append(actual)
                actual = []
        else:
            actual.append(el)
    if actual:
        tramos.append(actual)

    n = 0
    for tramo in tramos:
        n += 1
        ini = OxmlElement("w:permStart")
        ini.set(qn("w:id"), str(n))
        ini.set(qn("w:edGrp"), "everyone")
        fin = OxmlElement("w:permEnd")
        fin.set(qn("w:id"), str(n))
        tramo[0].addprevious(ini)
        tramo[-1].addnext(fin)

    _poner_documentProtection(doc, clave_)
    guardar_seguro(doc, ruta)
    if verbose:
        print(f"  Tramos de escritura libre: {n}")
        print(f"  Regiones de datos protegidas: "
              f"{sum(1 for h in hijos if _es_region_datos(h))}")
    return n
def desproteger(ruta, verbose=True):
    comprobar_escribible(ruta)
    doc = Document(str(ruta))
    settings = doc.settings.element
    quitados = 0
    for viejo in settings.findall(qn("w:documentProtection")):
        settings.remove(viejo)
        quitados += 1
    guardar_seguro(doc, ruta)
    if verbose:
        print("  Protección retirada." if quitados else "  No estaba protegido.")
    return quitados


# --------------------------------------------------------------------------- #
#  Rangos con nombre en el Excel
# --------------------------------------------------------------------------- #
def nombrar_rangos(xlsx, ctx, cfg, solo_simular=True, verbose=True):
    """Crea en el libro un nombre 'fs_<clave>' por cada línea del estado.

    Se hace con Excel (COM), NO con openpyxl: openpyxl no recalcula fórmulas
    y al reguardar descartaría el valor cacheado de todas ellas, dejando el
    Word en blanco. Excel guarda el libro con sus propias reglas y no toca
    ningún valor.

    Cada nombre apunta a la CELDA DE ETIQUETA de su fila. Así Excel reajusta
    la referencia solo cuando se insertan o borran filas encima, y el vínculo
    con el documento de Word no depende del texto de la etiqueta.
    """
    meta = ctx.get("_meta") or {}
    hoja = meta.get("hoja")
    col_etiqueta = (meta.get("columnas") or {}).get("etiqueta")
    if not hoja or not col_etiqueta or col_etiqueta == "—":
        raise ValueError(
            "No sé en qué hoja/columna poner los nombres.\n"
            "Ejecute primero 'verificar' para ver cómo se está leyendo el libro."
        )

    prefijo = str(cfg.get("prefijo_rangos") or "fs_")
    plan, vistas = [], set()
    for linea in ctx.get("lineas", []):
        fila = linea.get("fila")
        etiqueta = (linea.get("etiqueta") or "").strip()
        if not fila or not etiqueta:
            continue                              # subtotales sin etiqueta: se omiten
        k = linea.get("clave") or C.clave(etiqueta)
        if not k or k in vistas:
            continue
        vistas.add(k)
        plan.append({
            "nombre": f"{prefijo}{k}",
            "refiere": f"='{hoja}'!${col_etiqueta}${fila}",
            "fila": fila,
            "etiqueta": etiqueta,
            "ya": linea.get("clave_origen") == "rango",
        })

    nuevos = [p for p in plan if not p["ya"]]
    if verbose:
        print(f"  Líneas con etiqueta: {len(plan)}")
        print(f"  Ya tienen rango:     {len(plan) - len(nuevos)}")
        print(f"  Se crearían:         {len(nuevos)}")
    if solo_simular or not nuevos:
        return plan, 0

    return plan, _aplicar_nombres(xlsx, nuevos, verbose=verbose)


#: Conduce Excel desde PowerShell en vez de con pywin32. Evita una
#: dependencia binaria pesada (que además complica el empaquetado con
#: PyInstaller) y funciona en cualquier Windows con Excel instalado.
_PS_NOMBRAR = r"""
$ErrorActionPreference = 'Stop'
$plan  = Get-Content -Raw -Encoding UTF8 -LiteralPath $args[0] | ConvertFrom-Json
$libro = $args[1]
$excel = New-Object -ComObject Excel.Application
$excel.Visible = $false
$excel.DisplayAlerts = $false
$creados = 0
try {
  $wb = $excel.Workbooks.Open($libro)
  foreach ($p in $plan) {
    try { $wb.Names.Add($p.nombre, $p.refiere) | Out-Null; $creados++ }
    catch { Write-Output "  ! $($p.nombre): $($_.Exception.Message)" }
  }
  $wb.Save()
  $wb.Close($false)
} finally {
  $excel.Quit()
  [void][Runtime.InteropServices.Marshal]::ReleaseComObject($excel)
}
Write-Output "CREADOS=$creados"
"""


def _aplicar_nombres(xlsx, nuevos, verbose=True):
    import subprocess
    import tempfile

    tmp = Path(tempfile.mkdtemp(prefix="fs_nombrar_"))
    plan_json = tmp / "plan.json"
    plan_json.write_text(
        json.dumps([{"nombre": p["nombre"], "refiere": p["refiere"]} for p in nuevos],
                   ensure_ascii=False),
        encoding="utf-8",
    )

    try:
        stdout, stderr, codigo = ejecutar_ps(
            _PS_NOMBRAR, str(plan_json), str(Path(xlsx).resolve()),
            timeout=300, sta=False)
    except FileNotFoundError:
        raise ValueError(
            "No encontré PowerShell, que es lo que usa esta orden para pilotar Excel.\n"
            "Cree los nombres a mano: en Excel, seleccione la celda de la etiqueta\n"
            "y escriba el nombre en el Cuadro de nombres (arriba a la izquierda)."
        )
    except subprocess.TimeoutExpired:
        raise ValueError(
            "Excel tardó demasiado. ¿Está el libro abierto o pidiendo algo en pantalla?\n"
            "Ciérrelo y vuelva a intentarlo."
        )
    finally:
        shutil.rmtree(tmp, ignore_errors=True)

    salida = stdout + stderr
    creados = 0
    for linea in salida.splitlines():
        linea = linea.strip()
        if linea.startswith("CREADOS="):
            creados = int(linea.split("=", 1)[1])
        elif linea.startswith("!") or linea.startswith("! "):
            if verbose:
                print(f"    {linea}")
    if codigo != 0 and creados == 0:
        raise ValueError(
            "Excel no pudo escribir los nombres:\n"
            + "\n".join("  " + l for l in salida.strip().splitlines()[:12])
            + "\n\nCausa habitual: el libro está abierto en Excel. Ciérrelo."
        )
    if verbose:
        print(f"  Nombres creados: {creados}")
    return creados


# --------------------------------------------------------------------------- #
#  Columna «Tipo»: fijar lo que hoy se infiere
# --------------------------------------------------------------------------- #
#: Igual que _PS_NOMBRAR, se pilota Excel por COM. Con openpyxl NO se puede:
#: al reguardar descarta el valor cacheado de todas las formulas y el Word
#: saldria con las cifras en blanco.
_PS_TIPOS = r"""
$ErrorActionPreference = 'Stop'
$plan  = Get-Content -Raw -Encoding UTF8 -LiteralPath $args[0] | ConvertFrom-Json
$libro = $args[1]
$hoja  = $args[2]
$col   = [int]$args[3]
$excel = New-Object -ComObject Excel.Application
$excel.Visible = $false
$excel.DisplayAlerts = $false
$escritos = 0
try {
  $wb = $excel.Workbooks.Open($libro)
  $ws = $wb.Worksheets.Item($hoja)
  $ws.Cells.Item(1, $col).Value2 = 'Tipo'
  foreach ($p in $plan) {
    $ws.Cells.Item([int]$p.fila, $col).Value2 = [string]$p.tipo
    $escritos++
  }
  $ws.Columns.Item($col).ColumnWidth = 6
  $wb.Save()
  $wb.Close($false)
} finally {
  $excel.Quit()
  [void][Runtime.InteropServices.Marshal]::ReleaseComObject($excel)
}
Write-Output "ESCRITOS=$escritos"
"""


def fijar_tipos(xlsx, ctx, cfg, solo_simular=True, verbose=True):
    """Escribe en el libro la columna «Tipo» con lo que hoy se infiere.

    Mientras no exista esa columna, el tipo de cada fila se deduce de
    señales fragiles —negrita en las cifras, que la etiqueta empiece por
    'Total', que la fila no traiga numeros—. Funciona, pero cualquier
    retoque de formato puede cambiar en silencio como se clasifica una
    fila, y con ello el aspecto del documento.

    Fijarla convierte esa adivinanza en un dato declarado. No cambia nada
    hoy: escribe exactamente lo que ya se estaba infiriendo.

    Devuelve (plan, escritos).
    """
    meta = ctx.get("_meta") or {}
    hoja = meta.get("hoja")
    if not hoja:
        raise ValueError("No sé en qué hoja escribir la columna 'Tipo'.")
    if meta.get("hay_col_tipo"):
        if verbose:
            print("  El libro YA tiene columna 'Tipo': no se toca nada.")
        return [], 0

    col = _col_tipos_destino(cfg, meta)
    plan = [{"fila": l["fila"], "tipo": l.get("tipo", "I"),
             "etiqueta": (l.get("etiqueta") or "").strip()}
            for l in ctx.get("lineas", []) if l.get("fila")]

    if verbose:
        print(f"  Hoja:                {hoja}")
        print(f"  Columna de destino:  {get_column_letter(col)}  (libre, a la derecha)")
        print(f"  Filas a declarar:    {len(plan)}")
    if solo_simular or not plan:
        return plan, 0
    return plan, _aplicar_tipos(xlsx, plan, hoja, col, verbose=verbose)


def _col_tipos_destino(cfg, meta):
    """Columna donde va «Tipo»: la forzada en config.json, o la primera libre."""
    forzada = (cfg.get("columnas") or {}).get("tipo")
    if forzada:
        return column_index_from_string(str(forzada).strip().upper())
    usadas = [1]
    for letra in (meta.get("columnas") or {}).values():
        if letra and letra != "—":
            try:
                usadas.append(column_index_from_string(str(letra)))
            except ValueError:
                pass
    return max(usadas) + 1


def _aplicar_tipos(xlsx, plan, hoja, col, verbose=True):
    import subprocess
    import tempfile

    tmp = Path(tempfile.mkdtemp(prefix="fs_tipos_"))
    plan_json = tmp / "plan.json"
    plan_json.write_text(
        json.dumps([{"fila": p["fila"], "tipo": p["tipo"]} for p in plan],
                   ensure_ascii=False),
        encoding="utf-8")

    try:
        stdout, stderr, codigo = ejecutar_ps(
            _PS_TIPOS, str(plan_json), str(Path(xlsx).resolve()),
            str(hoja), str(col), timeout=300, sta=False)
    except FileNotFoundError:
        raise ValueError(
            "No encontré PowerShell, que es lo que usa esta orden para pilotar Excel.\n"
            f"Escriba a mano 'Tipo' en {get_column_letter(col)}1 y la letra de cada fila.")
    except subprocess.TimeoutExpired:
        raise ValueError(
            "Excel tardó demasiado. ¿Está el libro abierto? Ciérrelo y reintente.")
    finally:
        shutil.rmtree(tmp, ignore_errors=True)

    salida = stdout + stderr
    escritos = 0
    for linea in salida.splitlines():
        if linea.strip().startswith("ESCRITOS="):
            escritos = int(linea.split("=", 1)[1])
    if codigo != 0 and escritos == 0:
        raise ValueError(
            "Excel no pudo escribir la columna 'Tipo':\n"
            + "\n".join("  " + l for l in salida.strip().splitlines()[:12])
            + "\n\nCausa habitual: el libro está abierto en Excel. Ciérrelo.")
    if verbose:
        print(f"  Filas declaradas: {escritos}")
    return escritos


# --------------------------------------------------------------------------- #
#  Utilidades de línea de órdenes
# --------------------------------------------------------------------------- #
def _cargar_ctx(argumento=None):
    """Lee el Excel con el motor de generador_fs.

    Devuelve (ctx, xlsx, cfg). Se conserva ctx['_meta'] —la orden 'nombrar'
    necesita saber hoja y columna de etiqueta— y se descartan los avisos,
    que solo interesan al generador antiguo.
    """
    import generador_fs as G

    cfg = G.cargar_config()
    xlsx = Path(argumento).resolve() if argumento else G.encontrar_excel_por_convencion(cfg)
    if not xlsx.exists():
        raise ValueError(f"No se encontró el libro de Excel:\n  {xlsx}")
    ctx = G.leer_contexto(xlsx, cfg)
    ctx.pop("_avisos", None)
    return ctx, xlsx, cfg


def resolver_documento(argumento, cfg):
    """Decide sobre qué .docx se trabaja.

    Prioridad: lo que se indique en la orden; si no, config.json ->
    'documento_base' (ruta absoluta, o relativa a la raíz del proyecto).
    """
    if argumento:
        return Path(argumento).resolve()

    import generador_fs as G

    crudo = str(cfg.get("documento_base") or "").strip()
    if not crudo:
        raise ValueError(
            "No sé qué documento actualizar.\n\n"
            "Indíquelo en la orden:\n"
            "    ... refrescar \"MI_DOCUMENTO.docx\"\n\n"
            "O fíjelo de una vez en config.json:\n"
            '    "documento_base": "C:\\\\ruta\\\\a\\\\MI_DOCUMENTO.docx"'
        )
    # ${ONEDRIVE}, ${USUARIO}, ${PROYECTO}, ~ … Se resuelven aquí para que la
    # misma config.json valga en cualquier equipo.
    ruta = Path(G.expandir_ruta(crudo))
    if not ruta.is_absolute():
        ruta = (G.BASE / ruta).resolve()
    if ruta.exists():
        return ruta

    # Los nombres escritos en Office suelen colar espacios duros (U+00A0) y
    # acentos descompuestos, invisibles al ojo pero distintos byte a byte.
    # Antes de rendirnos, buscamos un archivo equivalente en la carpeta.
    cercano = _buscar_parecido(ruta) or _buscar_sin_tildes(ruta)
    if cercano is not None:
        print(f"(El nombre de config.json no coincidía exactamente; "
              f"uso '{cercano.name}')")
        return cercano

    # La ruta puede venir del perfil de OTRO equipo. config.json viaja entre
    # maquinas y, sobre todo, se queda EMBEBIDO en el .exe con la ruta de
    # quien lo compilo: en cuanto el .exe cambia de manos, ese
    # 'C:\\Users\\Fulano\\...' deja de existir y todas las opciones que
    # dependen del documento mueren a la vez. Antes de rendirse, se prueba
    # el mismo camino bajo el perfil de este usuario.
    reubicado = _reubicar_perfil(ruta)
    if reubicado is not None:
        print(f"(config.json apunta a un perfil que no existe en este equipo; "
              f"uso el equivalente aqui:\n    {reubicado})")
        return reubicado

    raise ValueError(
        f"config.json apunta a un documento que no existe:\n  {ruta}\n\n"
        "Revise la clave 'documento_base'. Ojo con los espacios: un nombre\n"
        "escrito en Office puede llevar espacios duros que no se ven.\n\n"
        "Para elegirlo con el explorador:\n"
        "    EstadosFinancieros.exe --elegir-documento"
    )


def _reubicar_perfil(ruta):
    """La misma ruta, pero bajo el perfil de usuario de ESTA maquina.

    Un 'C:\\Users\\Fulano\\...' que aqui no existe casi siempre es el mismo
    archivo bajo otro nombre de perfil: el equipo se renombro, o el .exe se
    compilo en otra maquina y se quedo con la ruta del que lo compilo.

    Se prueba, en orden:
      1. cambiar el tramo del perfil por el del usuario actual;
      2. lo mismo, tolerando espacios duros y acentos descompuestos;
      3. el mismo nombre de archivo en cualquier carpeta OneDrive del
         usuario actual.

    Devuelve None si nada encaja, y entonces el llamante avisa. Mejor un
    error claro que refrescar en silencio un documento que no era el que
    se pedia.
    """
    casa = Path.home()
    partes = ruta.parts
    # ('C:\\', 'Users', 'Fulano', 'OneDrive - X', 'documento.docx')
    if len(partes) > 3 and partes[1].lower() == "users":
        candidata = casa.joinpath(*partes[3:])
        if candidata.exists():
            return candidata
        cercana = _buscar_parecido(candidata)
        if cercana is not None:
            return cercana

    objetivo = _normalizar_nombre(ruta.name)
    for raiz in sorted(casa.glob("OneDrive*")):
        if not raiz.is_dir():
            continue
        for p in raiz.glob("*" + ruta.suffix):
            if _normalizar_nombre(p.name) == objetivo:
                return p
    return None


def _normalizar_nombre(nombre):
    """Colapsa cualquier tipo de espacio y unifica la forma Unicode."""
    import unicodedata

    s = unicodedata.normalize("NFC", str(nombre))
    return " ".join(s.split()).casefold()


def _buscar_parecido(ruta):
    """Un archivo de la misma carpeta cuyo nombre solo difiera en espacios
    o en la forma de los acentos. Devuelve None si no hay exactamente uno."""
    carpeta = ruta.parent
    if not carpeta.is_dir():
        return None
    objetivo = _normalizar_nombre(ruta.name)
    iguales = [p for p in carpeta.glob("*" + ruta.suffix)
               if _normalizar_nombre(p.name) == objetivo]
    return iguales[0] if len(iguales) == 1 else None


#: Lo que suele quedar donde habia un espacio duro (U+00A0) que una pagina
#: de codigos no supo escribir: la 'ÿ' de cp437/cp850, el '?' de una
#: sustitucion cualquiera, el caracter de reemplazo de Unicode.
_COMODINES = "ÿ?� "


def _sin_tildes(nombre):
    """El nombre sin diacriticos y sin distinguir tipos de espacio."""
    import unicodedata

    # Los comodines se cambian ANTES de normalizar: la 'ÿ' se descompone en
    # 'y' + diéresis, así que después de un NFKD ya no habría forma de
    # distinguirla de una 'y' de verdad.
    s = "".join(" " if c in _COMODINES else c for c in str(nombre))
    s = unicodedata.normalize("NFKD", s)
    s = "".join(c for c in s if not unicodedata.combining(c))
    return " ".join(s.split()).casefold()


def _buscar_sin_tildes(ruta):
    """Ultimo recurso: el mismo nombre ignorando tildes y espacios raros.

    Existe porque una ruta que ha pasado por la salida de una consola puede
    llegar con las tildes comidas ('EDICIÓN' -> 'EDICION') y el espacio duro
    convertido en otro byte. _buscar_parecido() no lo salva: normaliza la
    forma de los acentos, pero no los quita. Solo vale si hay UN candidato:
    ante la duda, mejor decir que no existe que refrescar el documento
    equivocado.
    """
    carpeta = ruta.parent
    if not carpeta.is_dir():
        return None
    objetivo = _sin_tildes(ruta.name)
    iguales = [p for p in carpeta.glob("*" + ruta.suffix)
               if _sin_tildes(p.name) == objetivo]
    return iguales[0] if len(iguales) == 1 else None


def _respaldar(ruta):
    """Copia previa, pero NUNCA sobre una copia buena con una mala.

    Un .docx es un ZIP. Si el archivo de partida no lo es (quedó a medio
    escribir, o Word lo tenía abierto), respaldarlo destruiría la única
    copia sana que queda. Mejor abortar y decirlo.
    """
    import zipfile

    ruta = Path(ruta)
    bak = ruta.with_suffix(ruta.suffix + ".bak")

    if not zipfile.is_zipfile(ruta):
        aviso = (
            f"El documento no es un .docx válido:\n  {ruta}\n\n"
            "Suele significar que se escribió sobre él mientras Word lo\n"
            "tenía abierto."
        )
        if bak.exists() and zipfile.is_zipfile(bak):
            aviso += (
                f"\n\nHay una copia previa SANA al lado:\n  {bak.name}\n"
                "Ciérrelo todo y restaure con:\n"
                f'    copy /Y "{bak.name}" "{ruta.name}"'
            )
        raise ValueError(aviso)

    shutil.copy2(ruta, bak)
    return bak


def _titulo(t):
    print()
    print("=" * 68)
    print(f" {t}")
    print("=" * 68)


# --------------------------------------------------------------------------- #
#  Elegir el documento de referencia
# --------------------------------------------------------------------------- #
#: Diálogo nativo de Windows. Se pilota desde PowerShell porque el Python
#: portable (distribución "embeddable") no trae tkinter, y añadirlo obligaría
#: a arrastrar Tcl/Tk dentro del .exe.
_PS_ELEGIR = r"""
$ErrorActionPreference = 'Stop'

# Nitidez en pantallas de alta densidad. Hay que declararlo ANTES de crear
# ninguna ventana, o Windows dibuja el dialogo a 96 ppp y lo escala como una
# imagen.
$codigo = @"
using System;
using System.Runtime.InteropServices;
public static class PppSel {
  [DllImport("user32.dll")] public static extern bool SetProcessDPIAware();
  [DllImport("user32.dll")] public static extern int SetProcessDpiAwarenessContext(IntPtr v);
}
"@
try { Add-Type -TypeDefinition $codigo -ErrorAction Stop } catch {}
try   { [void][PppSel]::SetProcessDpiAwarenessContext([IntPtr](-4)) }
catch { try { [void][PppSel]::SetProcessDPIAware() } catch {} }

Add-Type -AssemblyName System.Windows.Forms | Out-Null
Add-Type -AssemblyName System.Drawing | Out-Null
[System.Windows.Forms.Application]::EnableVisualStyles()

# $args[0] carpeta inicial   $args[1] archivo donde dejar la respuesta
$respuesta = if ($args.Count -ge 2) { $args[1] } else { '' }

$d = New-Object System.Windows.Forms.OpenFileDialog
$d.Title  = 'Elija el documento de Word que se actualizara'
$d.Filter = 'Documentos de Word (*.docx;*.docm;*.dotx;*.dotm)|*.docx;*.docm;*.dotx;*.dotm|Todos los archivos (*.*)|*.*'
$d.Multiselect = $false
$d.CheckFileExists = $true
if ($args.Count -ge 1 -and $args[0] -and (Test-Path -LiteralPath $args[0])) {
  $d.InitialDirectory = $args[0]
}

# Sin ventana duena, el dialogo sale DETRAS de la consola y sin foco: el
# usuario pulsa «Cambiar el documento», no ve nada, y el programa se queda
# esperando en silencio hasta que expira. Se le da como duena una ventana
# invisible y TopMost, que lo trae al frente.
$duena = New-Object System.Windows.Forms.Form
$duena.StartPosition   = 'CenterScreen'
$duena.Size            = New-Object System.Drawing.Size(1, 1)
$duena.FormBorderStyle = 'None'
$duena.ShowInTaskbar   = $false
$duena.Opacity         = 0
$duena.TopMost         = $true
$duena.Show()
$duena.Activate()
[System.Windows.Forms.Application]::DoEvents()

try {
  $r = $d.ShowDialog($duena)
} finally {
  $duena.Close()
  $duena.Dispose()
}

if ($r -eq [System.Windows.Forms.DialogResult]::OK) {
  # La ruta se devuelve por ARCHIVO, en UTF-8 sin BOM, no por la salida
  # estandar: ver _PS_UTF8. Un nombre con tildes o con espacios duros
  # —los que Word y OneDrive meten solos— no sobrevive a la pagina de
  # codigos de la consola, y el documento acababa rechazado por «No existe».
  [System.IO.File]::WriteAllText(
    $respuesta, $d.FileName, (New-Object System.Text.UTF8Encoding $false))
  Write-Output "ELEGIDO"
} else {
  Write-Output "CANCELADO"
}
"""


def _dialogo_ruta(script_ps, *argumentos, que="el explorador"):
    """Corre un dialogo de PowerShell que deja la ruta elegida en un archivo.

    Devuelve la ruta (Path), o None si se cancela o si el dialogo no llega a
    abrirse. El rodeo por archivo es deliberado: es la unica via que
    conserva el nombre byte a byte (ver _PS_UTF8).
    """
    import tempfile

    tmp = Path(tempfile.mkdtemp(prefix="fs_dialogo_"))
    respuesta = tmp / "ruta.txt"
    try:
        try:
            salida, error, _ = ejecutar_ps(
                script_ps, *argumentos, str(respuesta), timeout=600)
        except Exception as e:
            print(f"  No se pudo abrir {que} ({type(e).__name__}).")
            print("  Indique la ruta del documento a mano en config.local.json.")
            return None

        if respuesta.exists():
            crudo = respuesta.read_text(encoding="utf-8").strip()
            if crudo:
                return Path(crudo)
    finally:
        shutil.rmtree(tmp, ignore_errors=True)

    # Ni respuesta ni CANCELADO: el script se rompio. Callarselo deja al
    # usuario con un boton que aparenta no hacer nada.
    if "CANCELADO" not in salida:
        print(f"  {que.capitalize()} no llego a abrirse.")
        for l in (error or "").strip().splitlines()[:4]:
            print(f"    {l}")
    return None


def elegir_archivo_word(carpeta_inicial=None):
    """Abre el explorador de Windows filtrado a documentos de Word.

    Devuelve la ruta elegida, o None si se cancela o si no hay entorno
    gráfico (en cuyo caso el llamante debe pedirla por teclado).
    """
    return _dialogo_ruta(_PS_ELEGIR, str(carpeta_inicial or ""))


#: «Guardar como» para la plantilla que se crea desde cero. Es el mismo
#: dialogo de Windows de siempre, asi que el usuario elige indistintamente
#: una carpeta local o una de OneDrive: para el programa son la misma cosa.
_PS_GUARDAR = r"""
$ErrorActionPreference = 'Stop'

$codigo = @"
using System;
using System.Runtime.InteropServices;
public static class PppGuardar {
  [DllImport("user32.dll")] public static extern bool SetProcessDPIAware();
  [DllImport("user32.dll")] public static extern int SetProcessDpiAwarenessContext(IntPtr v);
}
"@
try { Add-Type -TypeDefinition $codigo -ErrorAction Stop } catch {}
try   { [void][PppGuardar]::SetProcessDpiAwarenessContext([IntPtr](-4)) }
catch { try { [void][PppGuardar]::SetProcessDPIAware() } catch {} }

Add-Type -AssemblyName System.Windows.Forms | Out-Null
Add-Type -AssemblyName System.Drawing | Out-Null
[System.Windows.Forms.Application]::EnableVisualStyles()

# $args[0] carpeta inicial   $args[1] nombre sugerido   $args[2] respuesta
$respuesta = if ($args.Count -ge 3) { $args[2] } else { '' }

$d = New-Object System.Windows.Forms.SaveFileDialog
$d.Title  = 'Donde quiere guardar la plantilla'
$d.Filter = 'Documento de Word (*.docx)|*.docx'
$d.DefaultExt = 'docx'
$d.AddExtension = $true
# Sin esto Windows pregunta «¿desea reemplazarlo?» y luego el programa
# preguntaria otra vez. La confirmacion la da el dialogo, una sola vez.
$d.OverwritePrompt = $true
if ($args.Count -ge 1 -and $args[0] -and (Test-Path -LiteralPath $args[0])) {
  $d.InitialDirectory = $args[0]
}
if ($args.Count -ge 2 -and $args[1]) { $d.FileName = $args[1] }

$duena = New-Object System.Windows.Forms.Form
$duena.StartPosition   = 'CenterScreen'
$duena.Size            = New-Object System.Drawing.Size(1, 1)
$duena.FormBorderStyle = 'None'
$duena.ShowInTaskbar   = $false
$duena.Opacity         = 0
$duena.TopMost         = $true
$duena.Show()
$duena.Activate()
[System.Windows.Forms.Application]::DoEvents()

try {
  $r = $d.ShowDialog($duena)
} finally {
  $duena.Close()
  $duena.Dispose()
}

if ($r -eq [System.Windows.Forms.DialogResult]::OK) {
  [System.IO.File]::WriteAllText(
    $respuesta, $d.FileName, (New-Object System.Text.UTF8Encoding $false))
  Write-Output "ELEGIDO"
} else {
  Write-Output "CANCELADO"
}
"""


def elegir_destino_word(carpeta_inicial=None, nombre_sugerido="Estados financieros.docx"):
    """«Guardar como» para el documento base que se va a crear.

    Devuelve la ruta elegida (aún inexistente) o None si se cancela. El
    usuario decide ahí si lo pone en el disco local o dentro de OneDrive;
    el programa trata las dos igual.
    """
    if carpeta_inicial is None:
        import generador_fs as G

        carpeta_inicial = G.raiz_onedrive() or Path.home()
    ruta = _dialogo_ruta(_PS_GUARDAR, str(carpeta_inicial or ""),
                         str(nombre_sugerido), que="la ventana de guardado")
    if ruta is not None and ruta.suffix.lower() != ".docx":
        ruta = ruta.with_suffix(".docx")
    return ruta


#: Las tres formas en que puede llegar un documento que el usuario elige.
#: Ninguna es un rechazo: las tres se saben trabajar.
LISTO = "listo"          # ya trae las regiones: se refresca y punto
EN_BLANCO = "en_blanco"  # vacio o casi: se usa tal cual como base
CON_TEXTO = "con_texto"  # tiene redaccion pero no regiones: se le añade un apartado

#: Cuantos parrafos con texto puede tener un documento y seguir contando
#: como «practicamente en blanco». Un Word recien creado trae uno o dos
#: parrafos vacios; uno que alguien empezo suele traer un titulo y poco mas.
_UMBRAL_EN_BLANCO = 3


def _contenido_visible(doc):
    """(parrafos con texto, tablas) FUERA de las regiones del contrato.

    Sirve para distinguir un documento en blanco de uno con redaccion. Lo
    que ya vive dentro de una region no cuenta: eso lo puso el programa.

    La pertenencia se mira subiendo por los padres, no comparando id(). En
    lxml el objeto de Python es un envoltorio que se crea al vuelo y se
    descarta en cuanto nadie lo referencia: dos nodos distintos pueden
    acabar en la misma direccion de memoria, y un conjunto de id() da falsos
    positivos. Con eso, un documento lleno de redaccion se contaba como
    vacio.
    """
    cuerpo = doc.element.body
    parrafos = 0
    for p in cuerpo.iter(qn("w:p")):
        if _dentro_de_region(p):
            continue
        if "".join(t.text or "" for t in p.iter(qn("w:t"))).strip():
            parrafos += 1
    tablas = sum(1 for t in cuerpo.iter(qn("w:tbl")) if not _dentro_de_region(t))
    return parrafos, tablas


def clasificar_documento(ruta):
    """¿Como hay que tratar este documento? Devuelve (estado, familias, detalle).

    `estado` es LISTO, EN_BLANCO o CON_TEXTO. Ninguno impide trabajar: es
    solo la diferencia entre refrescar, usar de base o añadir un apartado.
    """
    doc = Document(str(ruta))
    familias = {}
    for tag in _indexar(doc):
        fam, _, _ = C.descomponer(tag)
        familias[fam] = familias.get(fam, 0) + 1

    parrafos, tablas = _contenido_visible(doc)
    detalle = {"parrafos": parrafos, "tablas": tablas}

    if familias.get(C.FAM_TABLA):
        return LISTO, familias, detalle
    if parrafos <= _UMBRAL_EN_BLANCO and tablas == 0:
        return EN_BLANCO, familias, detalle
    return CON_TEXTO, familias, detalle


def revisar_candidato(ruta):
    """¿Sirve este archivo como documento base?

    Devuelve (ok, avisos, familias, info). `info` lleva 'ruta' —la real, que
    puede no ser la que se pidio— y 'estado' (LISTO / EN_BLANCO / CON_TEXTO).

    Solo se rechaza lo que de verdad no se puede abrir. Que le falten las
    regiones NO es motivo de rechazo: se le añaden. Antes esto era mas
    estricto de lo que hacia falta y dejaba fuera documentos utilizables.
    """
    import zipfile

    ruta = Path(ruta)
    avisos = []
    info = {"ruta": ruta, "estado": None}

    if not ruta.exists():
        # Un nombre escrito en Office trae espacios duros (U+00A0) y acentos
        # descompuestos que no se ven; y una ruta que ha pasado por una
        # consola puede haber perdido las tildes. Antes de rendirse, se
        # busca el archivo equivalente en la misma carpeta: es la diferencia
        # entre «no existe» y encontrarlo donde siempre estuvo.
        cercano = _buscar_parecido(ruta) or _buscar_sin_tildes(ruta)
        if cercano is None:
            return False, [f"No existe: {ruta}"], {}, info
        avisos.append(f"El nombre no coincidía exactamente; uso «{cercano.name}».")
        ruta = cercano
        info["ruta"] = ruta

    if ruta.suffix.lower() == ".doc":
        return False, avisos + [
            "Es un .doc del Word antiguo. Ábralo en Word y use «Guardar como»",
            "para dejarlo en .docx; ese sí sirve."], {}, info
    if ruta.suffix.lower() not in (".docx", ".docm", ".dotx", ".dotm"):
        avisos.append(f"La extensión '{ruta.suffix}' no es la de un documento de "
                      f"Word moderno. Puede fallar.")
    if not zipfile.is_zipfile(ruta):
        return False, avisos + [
            "No es un .docx válido. Si es un .doc antiguo, ábralo en Word y",
            "guárdelo como .docx; si no, puede que se quedara a medio escribir."], {}, info

    try:
        estado_, familias, detalle = clasificar_documento(ruta)
    except Exception as e:
        return False, avisos + [
            f"Word no lo puede abrir: {type(e).__name__}: {e}"], {}, info
    info["estado"] = estado_
    info.update(detalle)

    if estado_ == LISTO:
        avisos.append(
            f"Ya está integrado: tablas={familias.get(C.FAM_TABLA, 0)}  "
            f"campos={familias.get(C.FAM_CAMPO, 0)}  "
            f"cifras={familias.get(C.FAM_DATO, 0)}. Se refresca tal cual.")
    elif estado_ == EN_BLANCO:
        avisos.append("Está en blanco: se usará como base y se le montará "
                      "encima el estado completo.")
    else:
        avisos.append(
            f"Tiene redacción propia ({detalle['parrafos']} párrafos, "
            f"{detalle['tablas']} tablas) pero no las regiones del contrato.")
        avisos.append("Se le añadirá un apartado al final con las cifras; lo "
                      "que ya hay escrito no se toca.")

    culpables = quien_bloquea(ruta)
    if culpables:
        avisos.append("Está abierto ahora mismo: " + culpables[0])

    return True, avisos, familias, info


def fijar_documento_base(ruta_doc, verbose=True):
    """Guarda el documento elegido en la config de ESTE equipo.

    Va a config.local.json, no a config.json: el segundo viaja por git, y
    una ruta absoluta escrita ahí se le impone a la otra máquina en cada
    «pull». Además la ruta se guarda compactada (${ONEDRIVE}\\…) para que,
    si alguien la copia al config.json compartido, siga valiendo en ambas.
    """
    import generador_fs as G

    destino = G.CONFIG_LOCAL
    datos = {}
    if destino.exists():
        try:
            datos = json.loads(destino.read_text(encoding="utf-8"))
        except json.JSONDecodeError as e:
            raise ValueError(
                f"{destino.name} tiene un error de sintaxis y no me atrevo a\n"
                f"reescribirlo:\n  {e}"
            )
    datos.setdefault(
        "_comentario",
        "Ajustes de ESTE equipo. No se versiona: manda sobre config.json. "
        "Marcadores admitidos en las rutas: ${ONEDRIVE}, ${USUARIO}, ${PROYECTO}.")
    datos["documento_base"] = G.compactar_ruta(Path(ruta_doc).resolve())
    destino.write_text(
        json.dumps(datos, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    if verbose:
        print(f"  Guardado en {destino.name}  ->  {datos['documento_base']}")
    return destino
# --------------------------------------------------------------------------- #
#  Inspección del estado del proyecto
# --------------------------------------------------------------------------- #
def _marca(ok):
    return "[OK]  " if ok else "[--]  "


def _estado_documento(cfg):
    """Todo lo que se puede saber del documento base sin modificarlo."""
    import zipfile

    info = {"ruta": None, "existe": False, "valido": False, "bloqueado_por": [],
            "regiones": None, "protegido": None, "ultima": None, "origen": None,
            "candado": None, "error": None}
    try:
        ruta = resolver_documento(None, cfg)
    except ValueError as e:
        info["error"] = str(e).split("\n")[0]
        return info

    info["ruta"] = ruta
    info["existe"] = ruta.exists()
    if not info["existe"]:
        return info

    info["bloqueado_por"] = quien_bloquea(ruta)
    info["valido"] = zipfile.is_zipfile(ruta)
    if not info["valido"]:
        return info

    try:
        doc = Document(str(ruta))
        idx = _indexar(doc)
        familias = {}
        for tag in idx:
            fam, _, _ = C.descomponer(tag)
            familias[fam] = familias.get(fam, 0) + 1
        info["regiones"] = familias

        prot = doc.settings.element.find(qn("w:documentProtection"))
        info["protegido"] = prot.get(qn("w:edit")) if prot is not None else None
        info["candado"] = _contar_candados(doc)

        meta = _leer_meta(idx)
        info["ultima"] = meta.get("fecha")
        info["origen"] = meta.get("origen")
    except Exception as e:
        info["error"] = f"{type(e).__name__}: {e}"
    return info


def estado(cfg, xlsx_arg=None):
    """Radiografía del proyecto: los dos caminos, y qué falta en cada uno."""
    import generador_fs as G

    raiz = G.BASE
    _titulo("ESTADO DEL PROYECTO")

    # ---- 1. el documento vivo ---------------------------------------- #
    print()
    print(" 1. DOCUMENTO VIVO   (refrescar.bat / RefrescarFS.exe)")
    print(" " + "-" * 66)
    d = _estado_documento(cfg)
    if d["error"] and not d["existe"]:
        print(f"   {_marca(False)}Sin documento base configurado")
        print(f"         {d['error']}")
    else:
        print(f"   {_marca(d['existe'])}Documento: {d['ruta'].name}")
        print(f"         carpeta: {d['ruta'].parent}")
        if d["existe"]:
            print(f"   {_marca(d['valido'])}"
                  + ("Archivo íntegro (.docx válido)" if d["valido"]
                     else "ARCHIVO DAÑADO — no es un .docx válido"))
            if not d["valido"]:
                bak = d["ruta"].with_suffix(d["ruta"].suffix + ".bak")
                if bak.exists():
                    print(f"         hay copia previa al lado: {bak.name}")
        if d["bloqueado_por"]:
            print(f"   {_marca(False)}ABIERTO ahora mismo — no se puede refrescar")
            for c in d["bloqueado_por"]:
                print(f"         {c}")
        elif d["existe"]:
            print(f"   {_marca(True)}Libre (nadie lo tiene abierto)")

        if d["regiones"]:
            r = d["regiones"]
            preparado = r.get(C.FAM_TABLA, 0) > 0 and r.get(C.FAM_CAMPO, 0) > 0
            print(f"   {_marca(preparado)}"
                  + ("Preparado con sus regiones" if preparado
                     else "SIN preparar — ejecute: refrescar.bat --preparar"))
            print(f"         tablas={r.get(C.FAM_TABLA, 0)}  "
                  f"campos={r.get(C.FAM_CAMPO, 0)}  "
                  f"cifras en texto={r.get(C.FAM_DATO, 0)}  "
                  f"zonas de prosa={r.get(C.FAM_PROSA, 0)}")
            print(f"         bitácora={'sí' if r.get('registro') else 'NO'}  "
                  f"metadatos={'sí' if r.get('meta') else 'NO'}")
        if d["ultima"]:
            print(f"   {_marca(True)}Último refresco: {d['ultima']}")
            if d["origen"]:
                print(f"         origen: {d['origen']}")
        elif d["valido"]:
            print(f"   {_marca(False)}Nunca se ha refrescado")
        if d["candado"]:
            bloq, tot = d["candado"]
            if tot == 0:
                estado_cifras = "sin regiones de datos todavía"
            elif bloq == tot:
                estado_cifras = f"PROTEGIDAS ({bloq}/{tot} con candado)"
            elif bloq == 0:
                estado_cifras = f"EDITABLES A MANO (0/{tot} con candado)"
            else:
                estado_cifras = f"MIXTO ({bloq}/{tot} con candado)"
            print(f"   {_marca(True)}Cifras: {estado_cifras}")
        if d["protegido"]:
            print(f"   {_marca(True)}Modo estricto ACTIVO (protección: {d['protegido']})")
        elif d["valido"]:
            print(f"   {_marca(True)}Modo abierto (protección de documento: no)")

    # ---- 2. el camino clásico ---------------------------------------- #
    print()
    print(" 2. WORD NUEVO CADA VEZ   (generar.bat / GeneradorFS.exe)")
    print(" " + "-" * 66)
    plantilla = G.buscar_recurso(cfg["plantilla"])
    print(f"   {_marca(bool(plantilla))}Plantilla: "
          + (plantilla.name if plantilla else "NO ENCONTRADA"))
    salidas = raiz / "salidas"
    docs = (sorted(salidas.glob("*.docx"), key=lambda p: p.stat().st_mtime, reverse=True)
            if salidas.is_dir() else [])
    print(f"   {_marca(True)}Documentos generados en salidas\\: {len(docs)}")
    if docs:
        print(f"         más reciente: {docs[0].name}")

    # ---- 3. el libro de Excel ---------------------------------------- #
    print()
    print(" 3. LIBRO DE EXCEL")
    print(" " + "-" * 66)
    try:
        try:
            ctx, xlsx, _ = _cargar_ctx(xlsx_arg)
        except ValueError:
            # La convención de nombre no encontró nada: para una simple
            # inspección vale el libro de muestra de ejemplos\.
            muestras = sorted((raiz / "ejemplos").glob("*.xlsx"))
            if not muestras:
                raise
            ctx, xlsx, _ = _cargar_ctx(str(muestras[0]))
            print(f"   {_marca(True)}(sin libro propio; se inspecciona el de ejemplos\\)")
        meta = ctx.get("_meta", {})
        print(f"   {_marca(True)}Libro: {xlsx.name}")
        print(f"         hoja «{meta.get('hoja')}» ({meta.get('como_hoja')})")
        cols = meta.get("columnas", {})
        print(f"         columnas: etiqueta={cols.get('etiqueta')} "
              f"nota={cols.get('nota')} actual={cols.get('actual')} "
              f"previo={cols.get('previo')} tipo={cols.get('tipo')}")
        print(f"   {_marca(True)}Líneas: {len(ctx['lineas'])}  "
              f"({meta.get('n_declarados', 0)} declaradas, "
              f"{meta.get('n_inferidos', 0)} inferidas)")
        con_rango = meta.get("n_con_rango", 0)
        resto = len(ctx["lineas"]) - con_rango
        print(f"   {_marca(con_rango > 0)}Rangos con nombre: {con_rango} líneas")
        if resto > 0:
            print(f"         {resto} dependen aún del texto de la etiqueta")
            print("         para fijarlas: fs_documento.py nombrar <libro> --aplicar")
        if meta.get("n_escalares"):
            print(f"         escalares fuera de la tabla: {meta['n_escalares']}")

        if d.get("valido"):
            rep = verificar(d["ruta"], ctx)
            if rep["huerfanos"]:
                print(f"   {_marca(False)}Anclas huérfanas en el documento: "
                      f"{len(rep['huerfanos'])}")
                for t in rep["huerfanos"][:5]:
                    print(f"         ? {t}")
            else:
                print(f"   {_marca(True)}Sin anclas huérfanas")
            print(f"         cifras disponibles sin usar en la prosa: "
                  f"{len(rep['sin_usar'])}")
    except Exception as e:
        print(f"   {_marca(False)}No se pudo leer el Excel")
        print(f"         {type(e).__name__}: {str(e).splitlines()[0]}")

    # ---- 4. los ejecutables ------------------------------------------ #
    print()
    print(" 4. EJECUTABLES   (la vía local, sin complemento)")
    print(" " + "-" * 66)
    dist = raiz / "dist"
    esperados = ("GeneradorFS.exe", "RefrescarFS.exe", "EstadosFinancieros.exe")
    for nombre in esperados:
        exe = dist / nombre
        if exe.exists():
            mb = exe.stat().st_size / (1024 * 1024)
            fecha = datetime.fromtimestamp(exe.stat().st_mtime).strftime("%Y-%m-%d %H:%M")
            print(f"   {_marca(True)}{nombre}  ({mb:.1f} MB, {fecha})")
        else:
            print(f"   {_marca(False)}{nombre} sin generar")
    if not all((dist / n).exists() for n in esperados):
        print("         generarlos: tools\\hacer_exe.ps1")

    # ---- 5. el complemento ------------------------------------------- #
    print()
    print(" 5. COMPLEMENTO DE WORD")
    print(" " + "-" * 66)
    addin = raiz / "addin"
    print(f"   {_marca((addin / 'node_modules').is_dir())}"
          f"Dependencias instaladas (npm install)")
    print(f"   {_marca((addin / 'dist').is_dir())}Compilado (npm run build)")
    manifiesto = addin / "manifest.xml"
    if manifiesto.exists():
        texto_m = manifiesto.read_text(encoding="utf-8", errors="ignore")
        local = "localhost:3000" in texto_m
        print(f"   {_marca(not local)}"
              + ("Manifiesto apunta a una URL real" if not local
                 else "Manifiesto aún apunta a localhost:3000"))
    iconos = [addin / "assets" / f"icon-{n}.png" for n in (16, 32, 80)]
    print(f"   {_marca(all(i.exists() for i in iconos))}"
          f"Iconos ({sum(i.exists() for i in iconos)}/3)")
    print("         pasos completos en docs\\DESPLIEGUE_ADDIN.md")

    print()
    print("=" * 68)
    return 0


def main(argv):
    preparar_consola()
    if len(argv) < 2:
        print(__doc__)
        return 1

    orden = argv[1].lower()
    args = [a for a in argv[2:] if not a.startswith("--")]
    flags = {a.lower() for a in argv[2:] if a.startswith("--")}

    def opcion(nombre, defecto=None):
        for i, a in enumerate(argv):
            if a.lower() == nombre and i + 1 < len(argv):
                return argv[i + 1]
        return defecto

    if orden == "estado":
        import generador_fs as _G
        return estado(_G.cargar_config(), args[0] if args else None)

    if orden == "catalogo":
        ctx, xlsx, _cfg = _cargar_ctx(args[0] if args else None)
        _titulo(f"CIFRAS DISPONIBLES — {xlsx.name}")
        filas = C.catalogo(ctx)
        n_rango = sum(1 for f in filas if f[1] == "rango")
        print(f" {'clave':40} {'origen':9} {'actual':>16} {'previo':>16}")
        print(" " + "-" * 84)
        for k, origen, etiqueta, actual, previo in filas:
            print(f" {k:40} {origen:9} {actual:>16} {previo:>16}")
        print()
        print(f" {n_rango} de {len(filas)} claves vienen de un rango con nombre"
              f" (identidad estable).")
        if n_rango < len(filas):
            print(" Las de origen 'etiqueta' se rompen si alguien renombra la fila")
            print(" en el Excel. Para fijarlas:")
            print("     python fs_documento.py nombrar <libro.xlsx> --aplicar")
        print()
        print(" Para intercalar una cifra en la prosa, inserte en Word un control")
        print(" de contenido de TEXTO con la etiqueta (Tag):")
        print("     fs-dato-<clave>-actual     (o -previo, -nota, -var_abs, -var_pct)")
        print("=" * 68)
        return 0

    if orden == "tipos":
        ctx, xlsx, cfg = _cargar_ctx(args[0] if args else None)
        aplicar = "--aplicar" in flags
        _titulo(f"COLUMNA 'TIPO' — {xlsx.name}")
        if aplicar:
            bak = xlsx.with_suffix(xlsx.suffix + ".bak")
            shutil.copy2(xlsx, bak)
            print(f"  Copia previa del libro: {bak.name}")
        plan, escritos = fijar_tipos(xlsx, ctx, cfg, solo_simular=not aplicar)
        if plan and not aplicar:
            print()
            print("  Esto es una SIMULACIÓN. Para escribirla en el libro:")
            print(f"      python fs_documento.py tipos \"{xlsx.name}\" --aplicar")
            print()
            print("  Se escribe EXACTAMENTE lo que ya se venía infiriendo: el")
            print("  documento no cambia. Lo que cambia es que deja de depender")
            print("  de la negrita y del texto de la etiqueta.")
            print()
            resumen = {}
            for p in plan:
                resumen[p["tipo"]] = resumen.get(p["tipo"], 0) + 1
            print("  Reparto:  " + "   ".join(
                f"{t}={n}" for t, n in sorted(resumen.items())))
            print()
            for p in plan[:12]:
                print(f"     fila {p['fila']:>3}  {p['tipo']}   {p['etiqueta'][:44]}")
            if len(plan) > 12:
                print(f"       … y {len(plan) - 12} más.")
        print("=" * 68)
        return 0

    if orden == "nombrar":
        ctx, xlsx, cfg = _cargar_ctx(args[0] if args else None)
        aplicar = "--aplicar" in flags
        _titulo(f"RANGOS CON NOMBRE — {xlsx.name}")
        if aplicar:
            bak = xlsx.with_suffix(xlsx.suffix + ".bak")
            shutil.copy2(xlsx, bak)
            print(f"  Copia previa del libro: {bak.name}")
        plan, creados = nombrar_rangos(xlsx, ctx, cfg, solo_simular=not aplicar)
        if not aplicar:
            print()
            print("  Esto es una SIMULACIÓN. Para escribirlos en el libro:")
            print(f"      python fs_documento.py nombrar \"{xlsx.name}\" --aplicar")
            print()
            print("  Los nombres se crean con Excel, no con openpyxl: las fórmulas")
            print("  y sus valores cacheados quedan intactos.")
            print()
            print("  Primeros nombres del plan:")
            for p in plan[:12]:
                marca = "ya" if p["ya"] else " +"
                print(f"    {marca} {p['nombre']:44} -> fila {p['fila']}  ({p['etiqueta'][:28]})")
            if len(plan) > 12:
                print(f"       … y {len(plan) - 12} más.")
        print("=" * 68)
        return 0

    if orden == "plantilla":
        if not args:
            print("Falta el destino: plantilla <destino.docx>")
            return 1
        destino = Path(args[0]).resolve()
        _titulo(f"PLANTILLA BASE — {destino.name}")
        ctx, _cfg = None, None
        try:
            ctx, _, _cfg = _cargar_ctx(opcion("--excel"))
        except Exception:
            pass
        crear_base(destino, ctx or {}, _cfg)
        print(f" Escrita en: {destino}")
        print("=" * 68)
        return 0

    # De aquí en adelante todas las órdenes trabajan sobre un .docx. Si no
    # se indica, se toma el de config.json -> "documento_base".
    import generador_fs as _G

    doc_ruta = resolver_documento(args[0] if args else None, _G.cargar_config())
    if not doc_ruta.exists():
        print(f"No se encontró el documento:\n  {doc_ruta}")
        return 1
    if not args:
        args = [str(doc_ruta)]           # para que args[1] siga siendo el xlsx
        print(f"(Documento tomado de config.json: {doc_ruta.name})")

    if orden in ("construir", "reparar"):
        ctx = None
        try:
            ctx, _, _cfg = _cargar_ctx(args[1] if len(args) > 1 else None)
        except Exception:
            pass
        bak = _respaldar(doc_ruta)
        _titulo(f"ANDAMIAJE — {doc_ruta.name}")
        construir(doc_ruta, ctx or {}, cfg_bitacora=_cfg)
        print(f" Copia previa: {bak.name}")
        print("=" * 68)
        return 0

    if orden == "refrescar":
        ctx, xlsx, _cfg = _cargar_ctx(args[1] if len(args) > 1 else None)
        bak = _respaldar(doc_ruta)
        sha = hashlib.sha256(xlsx.read_bytes()).hexdigest()[:12]
        inf = refrescar(
            doc_ruta, ctx, origen=f"{xlsx.name} (sha {sha})",
            con_registro="--sin-registro" not in flags, cfg=_cfg,
        )
        _titulo(f"DOCUMENTO ACTUALIZADO — {doc_ruta.name}")
        print(f" Origen:        {xlsx.name}  sha256 {sha}")
        for nombre, n in inf["tablas"]:
            print(f" Tabla '{nombre}': {n} filas escritas")
        print(f" Campos de encabezado: {inf['campos']}")
        print(f" Cifras en prosa:      {inf['datos']}")
        print(f" Zonas de prosa intactas: {inf['sin_ancla_prosa']}")
        if inf.get("bitacora_archivo"):
            print(f" Bitácora:      {inf['bitacora_archivo']}")
        elif inf["con_registro"]:
            print(" Bitácora:      dentro del propio documento")
        else:
            print(" Bitácora:      (desactivada)")
        print(f" Copia previa: {bak.name}")
        if inf["huerfanos"]:
            print()
            print(" AVISO — anclas en el documento sin cifra en el Excel:")
            for t in inf["huerfanos"]:
                print(f"   ? {t}")
        if inf["colisiones"]:
            print()
            print(" AVISO — etiquetas distintas que producen la misma clave:")
            for nueva, primera, k in inf["colisiones"]:
                print(f"   '{nueva}' choca con '{primera}' (clave {k}); se usó la primera.")
        print()
        print(" Cambios respecto de la última actualización:")
        for c in inf["cambios"][:25]:
            print(imprimible(f"   * {c}"))
        if len(inf["cambios"]) > 25:
            print(imprimible(f"   • … y {len(inf['cambios']) - 25} más (ver la bitácora del documento)."))
        print("=" * 68)
        return 0

    if orden == "insertar":
        if len(args) < 3:
            print("Uso: insertar <doc.docx> <clave> <campo> "
                  "[--zona <nombre>] [--antes \"texto\"] [--despues \"texto\"]")
            print(f"     campos válidos: {', '.join(C.CAMPOS_DATO)}")
            print("     para ver las claves: python fs_documento.py catalogo")
            return 1
        bak = _respaldar(doc_ruta)
        _titulo(f"INSERTAR CIFRA — {doc_ruta.name}")
        insertar_dato(
            doc_ruta, args[1], args[2],
            zona=opcion("--zona", "analisis"),
            antes=opcion("--antes", ""),
            despues=opcion("--despues", ""),
        )
        print(f" Copia previa: {bak.name}")
        print(" Corra 'refrescar' para que tome el valor del Excel.")
        print("=" * 68)
        return 0

    if orden == "apariencia":
        modo = (args[1].lower() if len(args) > 1 else "")
        if modo not in ("visible", "invisible"):
            print("Uso: apariencia <doc.docx> <visible|invisible>")
            print()
            print("  visible    las cifras y la tabla llevan recuadro gris.")
            print("             Se ve de un vistazo qué mantiene el Excel.")
            print("  invisible  la cifra se lee como una palabra más del párrafo.")
            print("             Sigue bloqueada y sigue refrescándose igual.")
            return 1
        bak = _respaldar(doc_ruta)
        _titulo(f"APARIENCIA — {doc_ruta.name}")
        aspecto = "boundingBox" if modo == "visible" else "hidden"
        comprobar_escribible(doc_ruta)
        doc = Document(str(doc_ruta))
        n = normalizar_apariencia(doc, datos=aspecto)
        guardar_seguro(doc, doc_ruta)
        print(f"  Regiones ajustadas: {n}")
        print(f"  Copia previa: {bak.name}")
        print()
        if modo == "invisible":
            print(" Las cifras se leen ahora como texto normal del párrafo.")
            print(" Siguen bloqueadas: nadie puede teclear encima sin desbloquear.")
            print(" Para volver a verlas marcadas: apariencia <doc> visible")
        else:
            print(" Las cifras vuelven a mostrar su recuadro.")
        print("=" * 68)
        return 0

    if orden == "simplificar":
        bak = _respaldar(doc_ruta)
        _titulo(f"SIMPLIFICAR — {doc_ruta.name}")
        simplificar_documento(doc_ruta, quitar_prosa="--quitar-zonas" in flags)
        print(f" Copia previa: {bak.name}")
        print()
        print(" Los recuadros de las zonas de redacción y de los metadatos")
        print(" desaparecen. La tabla y las cifras conservan el suyo: es lo que")
        print(" deja ver de un vistazo qué mantiene el Excel.")
        if "--quitar-zonas" not in flags:
            print()
            print(" Para que las zonas de redacción dejen de existir del todo:")
            print("     fs_documento.py simplificar <doc.docx> --quitar-zonas")
        print("=" * 68)
        return 0

    if orden in ("desbloquear", "bloquear"):
        abrir = orden == "desbloquear"
        solo = args[1] if len(args) > 1 else None
        bak = _respaldar(doc_ruta)
        _titulo(("DESBLOQUEAR" if abrir else "BLOQUEAR") + f" — {doc_ruta.name}")
        cambiar_candado(doc_ruta, bloquear=not abrir, solo=solo)
        print(f" Copia previa: {bak.name}")
        if abrir:
            print()
            print(" Ya puede teclear encima de las cifras en Word.")
            print()
            print(" AVISO — lo que escriba a mano lo MACHACA el siguiente refresco:")
            print("         la región sigue vinculada al Excel. Si quiere que un")
            print("         valor escrito a mano sobreviva, desvincúlelo:")
            print("             fs_documento.py desvincular <doc> <clave>")
            print()
            print(" Para volver a proteger: fs_documento.py bloquear <doc>")
        else:
            print()
            print(" Las cifras vuelven a ser intocables a mano en Word.")
        print("=" * 68)
        return 0

    if orden == "desvincular":
        if len(args) < 2:
            print("Uso: desvincular <doc.docx> <clave-o-etiqueta>")
            print("     ejemplos:  total_assets        fs-dato-total_assets-actual")
            print("     ver las que hay:  fs_documento.py verificar <doc.docx>")
            return 1
        bak = _respaldar(doc_ruta)
        _titulo(f"DESVINCULAR — {doc_ruta.name}")
        quitadas = desvincular_region(doc_ruta, args[1])
        print(f" Copia previa: {bak.name}")
        if quitadas:
            print()
            print(" Ese texto ya NO se refresca: es prosa normal y se edita como tal.")
            print(" Para volver a vincularlo: fs_documento.py insertar <doc> <clave> <campo>")
        print("=" * 68)
        return 0

    if orden == "limpiar-bitacora":
        bak = _respaldar(doc_ruta)
        _titulo(f"LIMPIAR BITÁCORA — {doc_ruta.name}")
        quitar_registro_del_documento(doc_ruta)
        print(f" Copia previa: {bak.name}")
        print(" A partir de ahora el histórico va al .log de salidas\\.")
        print("=" * 68)
        return 0

    if orden == "verificar":
        ctx = None
        try:
            ctx, _, _cfg = _cargar_ctx(args[1] if len(args) > 1 else None)
        except Exception:
            pass
        rep = verificar(doc_ruta, ctx)
        _titulo(f"VERIFICACIÓN — {doc_ruta.name}")
        for nombre, tiene in rep["tablas"]:
            print(f" Tabla '{nombre}': {'con tabla dentro' if tiene else 'VACÍA (falta la tabla)'}")
        if not rep["tablas"]:
            print(" Tablas: NINGUNA. Corra 'construir'.")
        print(f" Campos de encabezado: {len(rep['campos'])}")
        for nombre, val in sorted(rep["campos"]):
            print(f"   {nombre:16} = {val!r}")
        print(f" Cifras en prosa: {len(rep['datos'])}")
        for nombre, campo, val in sorted(rep["datos"]):
            print(f"   {nombre}-{campo} = {val!r}")
        print(f" Zonas de prosa: {', '.join(sorted(rep['prosa'])) or '(ninguna)'}")
        print(f" Bitácora: {'sí' if rep['registro'] else 'NO'}    "
              f"Metadatos: {'sí' if rep['meta'] else 'NO'}")
        if rep["desconocidos"]:
            print(" Controles ajenos al contrato (se ignoran al refrescar):")
            for t in rep["desconocidos"]:
                print(f"   - {t}")
        if rep["huerfanos"]:
            print(" HUÉRFANOS — el documento los pide pero el Excel no los tiene:")
            for t in rep["huerfanos"]:
                print(f"   ? {t}")
        if rep["sin_usar"]:
            print(f" Cifras disponibles sin usar en la prosa: {len(rep['sin_usar'])}"
                  f"  (vea 'catalogo')")
        print("=" * 68)
        return 0

    if orden == "proteger":
        clave_ = opcion("--clave")
        if not clave_:
            print("Falta --clave <clave>. Es la que pedirá Word para desproteger.")
            return 1
        bak = _respaldar(doc_ruta)
        _titulo(f"PROTECCIÓN — {doc_ruta.name}")
        if "--salvo-datos" in flags:
            proteger_salvo_datos(doc_ruta, clave_)
        else:
            proteger(doc_ruta, clave_)
        print(f" Copia previa: {bak.name}")
        if "--salvo-datos" in flags:
            print(" Se puede escribir en TODO el documento MENOS en la tabla, los")
            print(" campos y las cifras. Esta protección la impone Word de verdad:")
            print(" ni Buscar y reemplazar ni Word en el navegador la saltan.")
        else:
            print(" Rol Redactor: solo puede escribir dentro de las zonas fs-prosa-*.")
        print("=" * 68)
        return 0

    if orden == "desproteger":
        bak = _respaldar(doc_ruta)
        _titulo(f"PROTECCIÓN — {doc_ruta.name}")
        desproteger(doc_ruta)
        print(f" Copia previa: {bak.name}")
        print("=" * 68)
        return 0

    print(f"Orden desconocida: {orden}")
    print(__doc__)
    return 1


if __name__ == "__main__":
    try:
        sys.exit(main(sys.argv))
    except ValueError as e:
        _titulo("NO SE PUDO COMPLETAR")
        print(str(e))
        print("=" * 68)
        sys.exit(1)
    except Exception:
        _titulo("ERROR INESPERADO — copie este texto para soporte")
        traceback.print_exc()
        print("=" * 68)
        sys.exit(1)
