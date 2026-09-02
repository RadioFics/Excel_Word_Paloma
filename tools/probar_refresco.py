"""
probar_refresco.py — Batería de pruebas del refresco Excel -> Word.

Responde a una pregunta concreta y repetible: si alguien añade una línea al
Excel, si la quita, si la renombra o si mete filas por encima, ¿el documento
de Word sale bien, y sigue siendo un .docx sano después de decenas de
actualizaciones seguidas?

Cada prueba trabaja sobre copias en una carpeta temporal. NUNCA toca el
documento real ni el libro real: se los pasa como plantilla de partida y
duplica lo que necesita.

Uso
---
    python tools\\probar_refresco.py                 todas las pruebas
    python tools\\probar_refresco.py --libro X.xlsx  además, una pasada
                                                     con un libro real
    python tools\\probar_refresco.py --verboso       enseña cada informe

Qué comprueba después de CADA escritura (esto es lo que detecta corrupción):

  * el archivo sigue siendo un ZIP válido y python-docx lo abre;
  * no hay Tags de región duplicados (dos regiones con el mismo nombre se
    pisarían la una a la otra en el siguiente refresco);
  * la región fs-meta sigue conteniendo JSON legible;
  * los párrafos de redacción siguen palabra por palabra como estaban;
  * el número de regiones no crece solo (una tabla que se duplica es el
    síntoma clásico de un refresco que inserta en vez de reemplazar).
"""
import argparse
import json
import shutil
import sys
import tempfile
import zipfile
from pathlib import Path

_RAIZ = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(_RAIZ / "src"))

from openpyxl import Workbook                       # noqa: E402
from docx import Document                           # noqa: E402

import fs_contrato as C                             # noqa: E402
import fs_documento as D                            # noqa: E402
import generador_fs as G                            # noqa: E402


ANCHO = 74
VERBOSO = False


# --------------------------------------------------------------------------- #
#  Un libro de Excel sintético
# --------------------------------------------------------------------------- #
#: Filas de partida: (etiqueta, nota, actual, previo, tipo). Es un estado de
#: situación financiera en miniatura, pero con la misma forma que el real:
#: encabezados, detalle, subtotales sin etiqueta y totales.
FILAS_BASE = [
    ("ASSETS",                     None, None,      None,      "H"),
    ("Current assets:",            None, None,      None,      "H"),
    ("Cash and cash equivalents",  None, 72_957_812, 81_370_000, "I"),
    ("Receivables",                1,    2_041_695,   813_700,  "I"),
    (None,                         None, 74_999_507, 82_183_700, "S"),
    ("Non-current assets:",        None, None,      None,      "H"),
    ("Mining concession asset",    1,    2_267_625,  6_621_077, "I"),
    ("Intangibles",                None,    48_822,      6_621, "I"),
    (None,                         None, 2_316_447,  6_627_698, "S"),
    ("Total assets",               None, 77_315_954, 88_811_398, "T"),
]


#: Cabecera del estado, tal como la escribe una persona en Excel:
#: fechas, escala, si está auditado y la moneda. Se guarda como
#: desplazamientos respecto de la esquina de la tabla, no como celdas fijas,
#: porque las pruebas la mueven por la hoja a propósito.
ENCABEZADO = [
    # (fila relativa, {desplazamiento de columna: valor})
    (0, {0: "As at", 2: "Note", 4: "June 5,", 5: "December 31,", 6: "Tipo"}),
    (1, {4: 1000, 5: 1000}),
    (2, {4: "(Unaudited)", 5: "(Audited)"}),
    (3, {4: "$", 5: "$"}),
]

#: Lo que debe salir del encabezado de arriba, se ponga donde se ponga.
ENCABEZADO_ESPERADO = {
    "titulo": "As at",
    "fecha_actual": "June 5,",
    "fecha_previa": "December 31,",
    "miles": 1000,
    "estado_actual": "Unaudited",
    "estado_previo": "Audited",
    "moneda": "$",
}


def escribir_libro(ruta, filas, origen=(1, 1)):
    """Crea un .xlsx con esas filas, con la tabla anclada en `origen`.

    `origen` es (fila, columna) de la esquina superior izquierda del
    encabezado. Por defecto A1, que es lo de siempre; las pruebas de
    posición la mueven a sitios como K14 para comprobar que el lector no
    depende de coordenadas fijas.

    Se usa openpyxl a propósito y sin miedo porque este libro lo generamos
    nosotros y no tiene ninguna fórmula que perder. Con el libro REAL nunca
    se hace: openpyxl descartaría los valores cacheados de las fórmulas.
    """
    r0, c0 = origen
    wb = Workbook()
    ws = wb.active
    ws.title = "FS"
    for dr, celdas in ENCABEZADO:
        for dc, valor in celdas.items():
            ws.cell(r0 + dr, c0 + dc).value = valor
    for i, (etq, nota, act, prev, tipo) in enumerate(filas):
        r = r0 + len(ENCABEZADO) + i
        ws.cell(r, c0 + 0).value = etq
        ws.cell(r, c0 + 2).value = nota
        ws.cell(r, c0 + 4).value = act
        ws.cell(r, c0 + 5).value = prev
        ws.cell(r, c0 + 6).value = tipo
    wb.save(ruta)
    return ruta


def poner_rangos(ruta, nombres):
    """Añade rangos con nombre al libro SINTÉTICO.

    Aquí sí se puede usar openpyxl: este libro lo hemos generado nosotros y
    no tiene ni una fórmula. Sobre el libro REAL hay que usar Excel por COM
    (fs_documento.py nombrar --aplicar), porque openpyxl al reguardar
    descarta los valores cacheados de todas las fórmulas.
    """
    from openpyxl import load_workbook
    from openpyxl.workbook.defined_name import DefinedName

    wb = load_workbook(ruta)
    for nombre, refiere in nombres.items():
        if nombre in wb.defined_names:
            del wb.defined_names[nombre]
        wb.defined_names.add(DefinedName(nombre, attr_text=refiere))
    wb.save(ruta)
    return ruta


def contexto_de(ruta, cfg):
    ctx = G.leer_contexto(ruta, cfg)
    ctx.pop("_avisos", None)
    return ctx


# --------------------------------------------------------------------------- #
#  Radiografía de un .docx, para comparar antes y después
# --------------------------------------------------------------------------- #
def _texto(el):
    """Todo el texto bajo un elemento, incluido lo que hay dentro de un sdt.

    Hace falta bajar al XML: doc.paragraphs y doc.tables de python-docx solo
    miran los hijos DIRECTOS del cuerpo, y aquí casi todo vive dentro de un
    control de contenido (w:sdt). Buscar con ellos da listas vacías y las
    comprobaciones pasan sin comprobar nada.
    """
    return "".join(t.text or "" for t in el.iter(D.qn("w:t")))


def _filas_de(el):
    """Las filas de tabla bajo `el`, como listas de textos de celda."""
    filas = []
    for tr in el.iter(D.qn("w:tr")):
        filas.append([_texto(tc) for tc in tr.iter(D.qn("w:tc"))])
    return filas


def radiografia(docx):
    """Todo lo que hace falta para detectar que algo se ha estropeado."""
    if not zipfile.is_zipfile(docx):
        raise AssertionError(f"{docx.name} ha dejado de ser un .docx válido (ZIP roto)")
    doc = Document(str(docx))
    cuerpo = doc.element.body

    tags, meta_crudo = [], None
    for sdt in cuerpo.iter(D.qn("w:sdt")):
        tag = D._tag_de(sdt)
        if tag:
            tags.append(tag)
        fam, _, _ = C.descomponer(tag)
        if fam == "meta":
            meta_crudo = _texto(sdt)

    return {
        "tags": tags,
        "tags_unicos": sorted(set(tags)),
        "n_regiones": len(tags),
        "filas": _filas_de(cuerpo),
        "filas_tabla": len(_filas_de(cuerpo)),
        "meta_crudo": meta_crudo,
        # La redacción son los párrafos de primer nivel del cuerpo: los que
        # una persona escribe. Lo de dentro de las regiones lo pone el Excel.
        "prosa": [p.text for p in doc.paragraphs],
        "texto": _texto(cuerpo),
        "bytes": docx.stat().st_size,
    }


def comprobar_sano(docx, etiqueta=""):
    """Las invariantes que deben cumplirse SIEMPRE, pase lo que pase."""
    r = radiografia(docx)
    fallos = []

    repetidos = [t for t in set(r["tags"]) if r["tags"].count(t) > 1]
    if repetidos:
        fallos.append(f"Tags de región duplicados: {repetidos}")

    if r["meta_crudo"] is not None:
        try:
            json.loads(r["meta_crudo"])
        except json.JSONDecodeError as e:
            fallos.append(f"fs-meta ya no es JSON legible: {e}")

    if not r["tags"]:
        fallos.append("El documento se ha quedado sin ninguna región")

    if fallos:
        raise AssertionError(f"[{etiqueta}] " + "; ".join(fallos))
    return r


# --------------------------------------------------------------------------- #
#  Andamiaje de una prueba
# --------------------------------------------------------------------------- #
class Banco:
    """Un documento y un libro recién hechos, listos para maltratarlos."""

    def __init__(self, carpeta, filas=None, prosa_extra=True):
        self.carpeta = carpeta
        self.cfg = G.cargar_config()
        self.cfg["bitacora"] = "no"          # sin .log: ensucia y no se mide aquí
        self.xlsx = escribir_libro(carpeta / "libro.xlsx", filas or FILAS_BASE)
        self.docx = carpeta / "documento.docx"
        # construir() AÑADE a un documento que ya existe (es lo que hace
        # «--preparar» sobre el Word de OneDrive). Aquí se parte de uno en
        # blanco, como quien empieza de cero.
        Document().save(str(self.docx))
        ctx = contexto_de(self.xlsx, self.cfg)
        D.construir(self.docx, ctx, verbose=False, cfg_bitacora=self.cfg)
        if prosa_extra:
            self._sembrar_prosa()

    def _sembrar_prosa(self):
        """Mete redacción humana en el documento.

        Es el corazón de la prueba: esta redacción NO puede cambiar nunca,
        por muchas veces que se refresque ni por mucho que cambie el Excel.
        """
        doc = Document(str(self.docx))
        doc.add_paragraph("Texto de Pamela que no se debe tocar jamás.")
        doc.add_paragraph("Segundo párrafo, con acentuación: gestión, análisis.")
        doc.save(str(self.docx))

    def refrescar(self, filas=None, origen="prueba"):
        if filas is not None:
            escribir_libro(self.xlsx, filas)
        ctx = contexto_de(self.xlsx, self.cfg)
        return D.refrescar(self.docx, ctx, origen=origen, verbose=False,
                           cfg=self.cfg)


# --------------------------------------------------------------------------- #
#  Las pruebas
# --------------------------------------------------------------------------- #
PRUEBAS = []


def prueba(titulo):
    def envoltorio(fn):
        PRUEBAS.append((titulo, fn))
        return fn
    return envoltorio


def _comprobar_encabezado(ctx, donde):
    """Fechas, escala, auditoría y moneda, salga la tabla donde salga."""
    fallos = []
    for campo, esperado in ENCABEZADO_ESPERADO.items():
        visto = ctx.get(campo, "")
        if str(esperado) not in str(visto):
            fallos.append(f"{campo}: esperaba ~{esperado!r} y hay {visto!r}")
    if fallos:
        raise AssertionError(f"[{donde}] encabezado mal leído -> " + "; ".join(fallos))


@prueba("El encabezado (fechas, moneda, escala) se lee con la tabla en A1")
def p_encabezado_en_origen(carpeta):
    cfg = G.cargar_config()
    xlsx = escribir_libro(carpeta / "libro.xlsx", FILAS_BASE)
    ctx = contexto_de(xlsx, cfg)
    _comprobar_encabezado(ctx, "A1")
    assert len(ctx["lineas"]) == len(FILAS_BASE), (
        f"esperaba {len(FILAS_BASE)} líneas y hay {len(ctx['lineas'])}: "
        "alguna fila del encabezado se coló como dato")
    return f"{len(ctx['lineas'])} líneas; encabezado completo"


@prueba("La MISMA tabla movida a K14 se lee exactamente igual")
def p_tabla_desplazada(carpeta):
    cfg = G.cargar_config()
    a = contexto_de(escribir_libro(carpeta / "a1.xlsx", FILAS_BASE), cfg)
    b = contexto_de(escribir_libro(carpeta / "k14.xlsx", FILAS_BASE,
                                   origen=(14, 11)), cfg)
    _comprobar_encabezado(b, "K14")

    assert len(b["lineas"]) == len(a["lineas"]), (
        f"en A1 salen {len(a['lineas'])} líneas y en K14 {len(b['lineas'])}")
    for campo in ENCABEZADO_ESPERADO:
        assert str(a.get(campo)) == str(b.get(campo)), (
            f"'{campo}' cambia al mover la tabla: "
            f"{a.get(campo)!r} vs {b.get(campo)!r}")
    for la, lb in zip(a["lineas"], b["lineas"]):
        assert la.get("etiqueta") == lb.get("etiqueta"), \
            f"etiquetas descuadradas: {la.get('etiqueta')!r} vs {lb.get('etiqueta')!r}"
        assert la.get("actual") == lb.get("actual"), \
            f"'{la.get('etiqueta')}': {la.get('actual')!r} vs {lb.get('actual')!r}"
        assert la.get("tipo") == lb.get("tipo"), \
            f"'{la.get('etiqueta')}': tipo {la.get('tipo')!r} vs {lb.get('tipo')!r}"
    return f"{len(b['lineas'])} líneas idénticas en A1 y en K14"


@prueba("La columna Tipo se detecta aunque quede lejos a la derecha")
def p_tipo_lejos(carpeta):
    cfg = G.cargar_config()
    ctx = contexto_de(escribir_libro(carpeta / "lejos.xlsx", FILAS_BASE,
                                     origen=(14, 11)), cfg)
    meta = ctx["_meta"]
    assert meta["hay_col_tipo"], (
        f"no encontró la columna 'Tipo' (quedó en la {meta['columnas']['tipo']}); "
        "con la tabla movida cae fuera del ancho de barrido")
    assert meta["n_declarados"] == len(FILAS_BASE), (
        f"solo {meta['n_declarados']} de {len(FILAS_BASE)} tipos vienen "
        f"declarados ({meta['n_inferidos']} se están infiriendo)")
    esperados = [f[4] for f in FILAS_BASE]
    vistos = [l.get("tipo") for l in ctx["lineas"]]
    assert vistos == esperados, f"tipos leídos {vistos} != declarados {esperados}"
    return (f"columna Tipo en {meta['columnas']['tipo']}, "
            f"{meta['n_declarados']} declarados")


@prueba("La hoja se encuentra por su ESTRUCTURA aunque la renombren")
def p_hoja_por_estructura(carpeta):
    from openpyxl import load_workbook

    cfg = G.cargar_config()
    xlsx = escribir_libro(carpeta / "renombrada.xlsx", FILAS_BASE, origen=(14, 11))

    wb = load_workbook(xlsx)
    wb["FS"].title = "Balance 2026 v3"          # ya no se llama como config
    señuelo = wb.create_sheet("FS", 0)          # y otra SÍ, pero vacía
    señuelo["A1"] = "notas sueltas"
    señuelo["A2"] = "esto no es un estado financiero"
    wb.save(xlsx)

    ctx = contexto_de(xlsx, cfg)
    meta = ctx["_meta"]
    assert meta["hoja"] == "Balance 2026 v3", (
        f"eligió la hoja '{meta['hoja']}'; debía elegir 'Balance 2026 v3' "
        f"por contenido, no el señuelo llamado 'FS'")
    assert len(ctx["lineas"]) == len(FILAS_BASE), (
        f"{len(ctx['lineas'])} líneas en vez de {len(FILAS_BASE)}")
    _comprobar_encabezado(ctx, "hoja renombrada")
    return f"eligió '{meta['hoja']}' ({meta['como_hoja']})"


@prueba("Refrescar dos veces seguidas no cambia nada la segunda vez")
def p_idempotencia(carpeta):
    b = Banco(carpeta)
    b.refrescar()
    antes = comprobar_sano(b.docx, "1er refresco")
    inf = b.refrescar()
    despues = comprobar_sano(b.docx, "2o refresco")

    assert despues["n_regiones"] == antes["n_regiones"], (
        f"las regiones pasaron de {antes['n_regiones']} a {despues['n_regiones']}")
    assert despues["filas_tabla"] == antes["filas_tabla"], (
        f"la tabla pasó de {antes['filas_tabla']} a {despues['filas_tabla']} filas")
    assert despues["prosa"] == antes["prosa"], "la redacción cambió sola"
    assert any("Sin cambios" in c for c in inf["cambios"]), (
        f"debería no reportar cambios, y reportó: {inf['cambios'][:3]}")
    return f"{antes['n_regiones']} regiones, {antes['filas_tabla']} filas, estables"


@prueba("Veinte refrescos seguidos no hinchan ni corrompen el documento")
def p_desgaste(carpeta):
    b = Banco(carpeta)
    b.refrescar()
    ref = comprobar_sano(b.docx, "referencia")
    for i in range(20):
        b.refrescar(origen=f"vuelta {i}")
        r = comprobar_sano(b.docx, f"vuelta {i}")
        assert r["n_regiones"] == ref["n_regiones"], (
            f"vuelta {i}: las regiones crecieron a {r['n_regiones']}")
        assert r["filas_tabla"] == ref["filas_tabla"], (
            f"vuelta {i}: la tabla creció a {r['filas_tabla']} filas")
        assert r["prosa"] == ref["prosa"], f"vuelta {i}: la redacción cambió"
    fin = radiografia(b.docx)
    crecimiento = fin["bytes"] / ref["bytes"]
    assert crecimiento < 1.5, f"el archivo creció x{crecimiento:.2f} en 20 refrescos"
    return f"20 vueltas; tamaño x{crecimiento:.2f}; {ref['filas_tabla']} filas siempre"


@prueba("Cambiar una cifra en el Excel cambia esa cifra y solo esa")
def p_cambio_de_cifra(carpeta):
    b = Banco(carpeta)
    b.refrescar()
    antes = comprobar_sano(b.docx, "antes")

    filas = [list(f) for f in FILAS_BASE]
    filas[2][2] = 99_999_999                      # Cash and cash equivalents
    inf = b.refrescar([tuple(f) for f in filas])
    despues = comprobar_sano(b.docx, "después")

    assert despues["filas_tabla"] == antes["filas_tabla"], "cambió el número de filas"
    assert despues["prosa"] == antes["prosa"], "la redacción cambió"
    texto = "\n".join(inf["cambios"])
    assert "99,999,999" in texto, f"el cambio no aparece en la bitácora: {inf['cambios']}"
    assert len(inf["cambios"]) == 1, (
        f"debería haber 1 cambio y hay {len(inf['cambios'])}: {inf['cambios']}")
    return f"1 cambio detectado: {inf['cambios'][0][:52]}"


@prueba("Añadir una línea al Excel la añade a la tabla del Word")
def p_anadir_linea(carpeta):
    b = Banco(carpeta)
    b.refrescar()
    antes = comprobar_sano(b.docx, "antes")

    filas = list(FILAS_BASE)
    filas.insert(8, ("VAT receivable", 10, 4_273_167, 662_108, "I"))
    b.refrescar(filas)
    despues = comprobar_sano(b.docx, "después")

    assert despues["filas_tabla"] == antes["filas_tabla"] + 1, (
        f"la tabla debería tener una fila más: {antes['filas_tabla']} -> "
        f"{despues['filas_tabla']}")
    assert despues["prosa"] == antes["prosa"], "la redacción cambió al añadir"
    etiquetas = [f[0] for f in despues["filas"] if f]
    assert "VAT receivable" in etiquetas, \
        f"la línea nueva no aparece en la tabla; hay: {etiquetas}"
    fila = next(f for f in despues["filas"] if f and f[0] == "VAT receivable")
    assert "4,273,167" in " ".join(fila), f"la línea nueva sin su cifra: {fila}"
    return f"{antes['filas_tabla']} -> {despues['filas_tabla']} filas, prosa intacta"


@prueba("Quitar una línea del Excel la quita de la tabla del Word")
def p_quitar_linea(carpeta):
    b = Banco(carpeta)
    b.refrescar()
    antes = comprobar_sano(b.docx, "antes")

    filas = [f for f in FILAS_BASE if f[0] != "Intangibles"]
    b.refrescar(filas)
    despues = comprobar_sano(b.docx, "después")

    assert despues["filas_tabla"] == antes["filas_tabla"] - 1, (
        f"la tabla debería tener una fila menos: {antes['filas_tabla']} -> "
        f"{despues['filas_tabla']}")
    assert despues["prosa"] == antes["prosa"], "la redacción cambió al quitar"
    etiquetas = [f[0] for f in despues["filas"] if f]
    assert "Intangibles" in [f[0] for f in antes["filas"] if f], \
        "la prueba no vale: 'Intangibles' no estaba antes de quitarla"
    assert "Intangibles" not in etiquetas, \
        f"la línea borrada sigue en la tabla; hay: {etiquetas}"
    return f"{antes['filas_tabla']} -> {despues['filas_tabla']} filas, prosa intacta"


@prueba("Una cifra intercalada en la redacción se actualiza sola")
def p_cifra_en_prosa(carpeta):
    b = Banco(carpeta)
    b.refrescar()
    D.insertar_dato(b.docx, "total_assets", "actual", zona="introduccion",
                    antes="Los activos totales ascendieron a ", despues=".",
                    verbose=False)
    b.refrescar()
    r = comprobar_sano(b.docx, "tras insertar")

    assert "Los activos totales ascendieron a " in r["texto"], \
        "no se insertó el texto de acompañamiento en la zona de redacción"
    assert "77,315,954" in r["texto"], \
        f"la cifra no se escribió en la prosa: {r['texto'][:300]}"

    filas = [list(f) for f in FILAS_BASE]
    filas[-1][2] = 12_345_678                     # Total assets
    b.refrescar([tuple(f) for f in filas])
    r = comprobar_sano(b.docx, "tras cambiar")
    assert "12,345,678" in r["texto"], "la cifra de la prosa no siguió al Excel"
    assert "77,315,954" not in r["texto"], "quedó la cifra vieja en la prosa"
    return "la cifra dentro del párrafo siguió al Excel"


@prueba("Si el Excel pierde una fila que la redacción cita, se avisa")
def p_huerfano(carpeta):
    b = Banco(carpeta)
    b.refrescar()
    D.insertar_dato(b.docx, "intangibles", "actual", zona="introduccion",
                    antes="Intangibles: ", despues=".", verbose=False)
    b.refrescar()

    filas = [f for f in FILAS_BASE if f[0] != "Intangibles"]
    inf = b.refrescar(filas)
    comprobar_sano(b.docx, "con huérfano")

    assert inf["huerfanos"], (
        "quitar del Excel una fila citada en la redacción debería avisar")
    assert any("intangibles" in h for h in inf["huerfanos"]), \
        f"el aviso no menciona la fila: {inf['huerfanos']}"
    return f"avisó de {len(inf['huerfanos'])} ancla(s) huérfana(s)"


@prueba("Renombrar una etiqueta rompe el vínculo si no hay rango con nombre")
def p_renombrar_sin_rango(carpeta):
    b = Banco(carpeta)
    b.refrescar()
    D.insertar_dato(b.docx, "intangibles", "actual", zona="introduccion",
                    antes="Intangibles: ", despues=".", verbose=False)
    b.refrescar()

    filas = [list(f) for f in FILAS_BASE]
    filas[7][0] = "Intangible assets"             # mismo concepto, otro texto
    inf = b.refrescar([tuple(f) for f in filas])
    comprobar_sano(b.docx, "tras renombrar")

    assert inf["huerfanos"], (
        "sin rango con nombre, renombrar la etiqueta DEBE dejar huérfana el "
        "ancla: es justo la fragilidad que documentamos")
    return f"confirmada la fragilidad conocida: {inf['huerfanos'][0]}"


@prueba("Con rango con nombre, renombrar la etiqueta NO rompe el vínculo")
def p_renombrar_con_rango(carpeta):
    b = Banco(carpeta)
    # 'Intangibles' está en la fila 12 del libro sintético (5 + índice 7).
    poner_rangos(b.xlsx, {"fs_intangibles": "'FS'!$A$12"})
    b.refrescar()
    D.insertar_dato(b.docx, "intangibles", "actual", zona="introduccion",
                    antes="Intangibles: ", despues=".", verbose=False)
    b.refrescar()
    r = comprobar_sano(b.docx, "con rango")
    assert "48,822" in r["texto"], "la cifra no llegó a la prosa"

    filas = [list(f) for f in FILAS_BASE]
    filas[7][0] = "Intangible assets"             # mismo concepto, otro texto
    filas[7][2] = 55_555                          # y otra cifra
    escribir_libro(b.xlsx, [tuple(f) for f in filas])
    poner_rangos(b.xlsx, {"fs_intangibles": "'FS'!$A$12"})
    inf = b.refrescar()
    r = comprobar_sano(b.docx, "renombrado con rango")

    assert not inf["huerfanos"], (
        "con rango con nombre NO debería quedar huérfana ninguna ancla, y "
        f"quedaron: {inf['huerfanos']}")
    assert "55,555" in r["texto"], \
        "la cifra de la prosa no siguió a la fila renombrada"
    return "el vínculo sobrevivió al renombrado (esto es lo que arregla el rango)"


@prueba("Cambiar el orden de las filas no descoloca las cifras")
def p_reordenar(carpeta):
    b = Banco(carpeta)
    b.refrescar()
    antes = comprobar_sano(b.docx, "antes")

    filas = list(FILAS_BASE)
    filas[6], filas[7] = filas[7], filas[6]       # Intangibles <-> Mining concession
    b.refrescar(filas)
    despues = comprobar_sano(b.docx, "después")

    assert despues["filas_tabla"] == antes["filas_tabla"], "cambió el número de filas"
    esperado = {"Intangibles": "48,822", "Mining concession asset": "2,267,625"}
    vistas = {}
    for f in despues["filas"]:
        if f and f[0].strip() in esperado:
            vistas[f[0].strip()] = " ".join(f[1:])
    assert set(vistas) == set(esperado), (
        f"no encontré las dos filas reordenadas; vi: {sorted(vistas)}")
    for etq, cifra in esperado.items():
        assert cifra in vistas[etq], (
            f"tras reordenar, '{etq}' muestra '{vistas[etq]}' y debería "
            f"contener '{cifra}'")
    return f"cada cifra siguió a su etiqueta ({len(vistas)} comprobadas)"


@prueba("Un documento con la redacción borrada sigue refrescándose")
def p_prosa_borrada(carpeta):
    b = Banco(carpeta)
    b.refrescar()
    doc = Document(str(b.docx))
    for p in list(doc.paragraphs):
        if p.text.startswith("Texto de Pamela"):
            p._element.getparent().remove(p._element)
    doc.save(str(b.docx))

    b.refrescar()
    r = comprobar_sano(b.docx, "sin prosa")
    assert not any(p.startswith("Texto de Pamela") for p in r["prosa"]), \
        "el refresco resucitó un párrafo que se había borrado"
    return "lo borrado sigue borrado; el refresco no lo repone"


@prueba("El respaldo .bak se crea y sirve para volver atrás")
def p_respaldo(carpeta):
    b = Banco(carpeta)
    b.refrescar()
    antes = radiografia(b.docx)

    D._respaldar(b.docx)
    bak = b.docx.with_suffix(b.docx.suffix + ".bak")
    assert bak.exists(), "no se creó la copia .bak"

    filas = [list(f) for f in FILAS_BASE]
    filas[2][2] = 1                               # destroza una cifra
    b.refrescar([tuple(f) for f in filas])
    assert radiografia(b.docx)["bytes"] != 0

    shutil.copy2(bak, b.docx)                     # restaurar
    vuelto = comprobar_sano(b.docx, "restaurado")
    assert vuelto["filas_tabla"] == antes["filas_tabla"], \
        "el documento restaurado no coincide con el original"
    return "respaldo creado y restauración correcta"


@prueba("Un .docx corrupto se rechaza en vez de destruir la copia buena")
def p_docx_corrupto(carpeta):
    b = Banco(carpeta)
    b.refrescar()
    D._respaldar(b.docx)

    b.docx.write_bytes(b"esto no es un zip")      # simula escritura a medias
    try:
        D._respaldar(b.docx)
    except ValueError as e:
        assert "no es un .docx" in str(e), f"mensaje inesperado: {e}"
        assert ".bak" in str(e), "no mencionó la copia sana que hay al lado"
        return "rechazado, y avisa de la copia sana"
    raise AssertionError("respaldó un archivo corrupto encima de la copia buena")


# --------------------------------------------------------------------------- #
#  Pasada con el libro real
# --------------------------------------------------------------------------- #
def prueba_libro_real(ruta_libro, carpeta):
    """Refresca dos veces con el libro de verdad, sobre un documento nuevo.

    No modifica el libro: solo lo lee.
    """
    cfg = G.cargar_config()
    cfg["bitacora"] = "no"
    ctx = contexto_de(Path(ruta_libro), cfg)
    docx = carpeta / "real.docx"
    Document().save(str(docx))
    D.construir(docx, ctx, verbose=False, cfg_bitacora=cfg)

    doc = Document(str(docx))
    doc.add_paragraph("Redacción de prueba sobre el libro real.")
    doc.save(str(docx))

    D.refrescar(docx, ctx, origen="real 1", verbose=False, cfg=cfg)
    a = comprobar_sano(docx, "real 1")
    inf = D.refrescar(docx, ctx, origen="real 2", verbose=False, cfg=cfg)
    b = comprobar_sano(docx, "real 2")

    assert a["filas_tabla"] == b["filas_tabla"], "la tabla cambió en el 2o refresco"
    assert a["prosa"] == b["prosa"], "la redacción cambió en el 2o refresco"
    assert any("Sin cambios" in c for c in inf["cambios"]), \
        f"el 2o refresco debería no ver cambios: {inf['cambios'][:3]}"

    meta = ctx.get("_meta") or {}
    return (f"hoja '{meta.get('hoja')}', {len(ctx['lineas'])} líneas, "
            f"{a['filas_tabla']} filas de tabla, {a['n_regiones']} regiones")


# --------------------------------------------------------------------------- #
#  Ejecución
# --------------------------------------------------------------------------- #
def main():
    global VERBOSO
    ap = argparse.ArgumentParser(description="Pruebas del refresco Excel -> Word")
    ap.add_argument("--libro", help="Un .xlsx real para una pasada adicional")
    ap.add_argument("--verboso", action="store_true")
    args = ap.parse_args()
    VERBOSO = args.verboso

    D.preparar_consola()
    print("=" * ANCHO)
    print(" PRUEBAS DEL REFRESCO   Excel -> Word")
    print("=" * ANCHO)
    print(" Cada prueba usa copias en una carpeta temporal.")
    print(" El documento y el libro reales no se tocan.")
    print()

    raiz = Path(tempfile.mkdtemp(prefix="fs_pruebas_"))
    ok = fallos = 0
    try:
        for i, (titulo, fn) in enumerate(PRUEBAS, 1):
            carpeta = raiz / f"c{i:02}"
            carpeta.mkdir(parents=True, exist_ok=True)
            print(f" {i:2}. {titulo}")
            try:
                detalle = fn(carpeta)
                ok += 1
                print(f"     BIEN   {detalle}")
            except AssertionError as e:
                fallos += 1
                print(f"     FALLO  {e}")
            except Exception as e:
                fallos += 1
                print(f"     ERROR  {type(e).__name__}: {e}")
                if VERBOSO:
                    import traceback
                    traceback.print_exc()
            print()

        if args.libro:
            carpeta = raiz / "real"
            carpeta.mkdir(parents=True, exist_ok=True)
            print(f" {len(PRUEBAS) + 1:2}. Pasada con el libro real "
                  f"({Path(args.libro).name})")
            try:
                print(f"     BIEN   {prueba_libro_real(args.libro, carpeta)}")
                ok += 1
            except AssertionError as e:
                fallos += 1
                print(f"     FALLO  {e}")
            except Exception as e:
                fallos += 1
                print(f"     ERROR  {type(e).__name__}: {e}")
                if VERBOSO:
                    import traceback
                    traceback.print_exc()
            print()
    finally:
        shutil.rmtree(raiz, ignore_errors=True)

    print("=" * ANCHO)
    print(f" {ok} bien, {fallos} mal")
    print("=" * ANCHO)
    return 1 if fallos else 0


if __name__ == "__main__":
    sys.exit(main())
